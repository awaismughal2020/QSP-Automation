"""
Word Template Updater

Updates the quarterly report Word template with:
- Global find/replace for quarter references
- Specific value updates on designated pages
- Date updates
- PAGE 4: Numeric KPI population (GTRI, gross rental income, vacancy, maintenance, etc.)
- PAGE 5-6: Actions and CAPEX values

NUMERIC POPULATION (Page 4):
The Word document contains financial KPIs embedded in narrative text like:
"GTRI of the portfolio amounted to €3,200k"
"Financial vacancy was 5.4%"
etc.

These values must be extracted from the generated Excel outputs and injected
into the Word document, replacing the previous quarter's values.

VALUE SOURCES (from Management Cijfers sheet):
- Column AA: Current quarter values (Q3 2025 = column 27)
- Column AB: LTM values (Q3 2025 LTM = column 28)
- Dynamic column calculation: base_column + (quarter_total - base_quarter_total)
"""

from dataclasses import dataclass, field
from typing import Dict, List, Optional, Tuple, Any
from pathlib import Path
from datetime import datetime
import subprocess
import re
import shutil
import tempfile
from loguru import logger


@dataclass
class ReportValues:
    """Values to insert into the quarterly report."""
    # Page 1
    report_date: str  # e.g., "14 October 2025"
    
    # Page 4 - Executive Summary (from Management Accounts / BDO)
    gtri: float  # Gross Theoretical Rental Income (€k) - quarterly
    gtri_ltm: float = 0.0  # GTRI LTM (€k) - for annual references / rent roll yields
    gross_rental_income: float = 0.0  # Actual rental income (€k) - quarterly
    gross_rental_income_ltm: float = 0.0  # Gross rental income LTM (€k)
    financial_vacancy_pct: float = 0.0  # Vacancy percentage
    financial_vacancy_amount: float = 0.0  # Vacancy in €k
    
    # Page 4 - From Rent Roll
    rent_roll_annual: float = 0.0  # Annual rent roll total (€k)
    rent_roll_units: int = 0  # Number of units in rent roll
    
    # Page 4 - From Sales Tracker
    units_sold_quarter: int = 0  # Units sold this quarter
    unit_sales_proceeds: float = 0.0  # Proceeds from sales (€k)
    
    # Page 4/6 - Maintenance and CAPEX
    maintenance_amount: float = 0.0  # Maintenance spend (€k) - quarterly
    maintenance_ltm: float = 0.0  # Maintenance LTM (€k)
    capex_amount: float = 0.0  # CAPEX spend (€k) - quarterly
    capex_ltm: float = 0.0  # CAPEX LTM (€k)
    
    # Page 5 - Unit Sales narrative
    unit_sales_narrative: str = ""  # Free text about unit sales actions
    
    # Page 6
    maintenance_detail: str = ""  # Maintenance description
    sustainability_detail: str = ""  # Sustainability/CAPEX description
    
    # Previous quarter values (for finding/replacing)
    # These are extracted from the template document (Q2 values)
    prev_gtri: float = 0.0
    prev_gtri_ltm: float = 0.0  # Previous LTM GTRI (for rent roll yields)
    prev_gross_rental_income: float = 0.0
    prev_vacancy_pct: float = 0.0
    prev_vacancy_amount: float = 0.0
    prev_rent_roll: float = 0.0
    prev_maintenance: float = 0.0
    prev_unit_sales_proceeds: float = 0.0


def extract_values_from_management_accounts(ma_path: str, quarter: int, year: int) -> dict:
    """
    Extract ALL financial values from Management Accounts Excel file.
    
    ALL values are extracted from Management Cijfers sheet using FULLY DYNAMIC column selection.
    The column is automatically calculated based on quarter and year - NO HARDCODING.
    
    Column Calculation Formula (fully dynamic):
        quarterly_column = 27 + ((year * 4 + quarter) - (2025 * 4 + 3))
    
    Examples (automatically calculated):
        Q3 2025 → Column AA (index 27)  [base reference]
        Q4 2025 → Column AB (index 28)  [auto-calculated]
        Q1 2026 → Column AC (index 29)  [auto-calculated]
        Q2 2026 → Column AD (index 30)  [auto-calculated]
        ... and so on for ALL future quarters
    
    Row mappings in Management Cijfers (rows are fixed, columns are dynamic):
        - Row 23: Gross Theoretical rental income (GTRI)
        - Row 24: (Financial vacancy) amount
        - Row 25: Gross rental income (calculated: GTRI - vacancy)
        - Row 26: Vacancy % (calculated)
        - Row 32: (Maintenance & repair costs actual)
        - Row 50: Cash proceeds sale (unit sale proceeds)
    
    LTM (Last Twelve Months) values are extracted from the column immediately
    to the right of the quarterly column (quarterly_column + 1).
    
    Args:
        ma_path: Path to Management Accounts Excel file
        quarter: Quarter number (1-4)
        year: Year (e.g., 2025)
        
    Returns:
        Dictionary with extracted values in €k
        
    Note: This function is FULLY DYNAMIC and will automatically work for Q4 2025,
    Q1 2026, and all future quarters without any code changes. The base reference
    (Q3 2025 = Column AA) is only used as a calculation anchor point.
    """
    import openpyxl
    from openpyxl.utils import get_column_letter
    
    try:
        # Load with data_only=True to get calculated values (formulas evaluated)
        wb = openpyxl.load_workbook(ma_path, data_only=True)
        
        # Find Management Cijfers sheet
        cijfers_sheet = None
        for name in wb.sheetnames:
            if 'Management Cijfers' in name:
                cijfers_sheet = wb[name]
                logger.info(f"Found Management Cijfers sheet: {name}")
                break
        
        if cijfers_sheet is None:
            logger.warning("Management Cijfers sheet not found")
            wb.close()
            return {}
        
        from ..transformers.management_accounts import (
            find_ltm_column_near_quarter,
            resolve_quarter_column,
        )

        header_row = 22
        quarter_label = f"Q{quarter} {year}"
        quarterly_column = resolve_quarter_column(
            cijfers_sheet, quarter_label, header_row
        )
        ltm_column = None
        if quarterly_column:
            ltm_column = find_ltm_column_near_quarter(
                cijfers_sheet, quarterly_column, header_row
            )

        if not quarterly_column:
            # Fallback: index from Q3 2025 = column AA (27) when headers are missing
            base_quarter_total = 2025 * 4 + 3
            target_quarter_total = year * 4 + quarter
            quarterly_column = 27 + (target_quarter_total - base_quarter_total)
            ltm_column = quarterly_column + 1
            logger.warning(
                f"Q{quarter} {year}: header {quarter_label!r} not found; "
                f"using column index {quarterly_column}"
            )

        if not ltm_column:
            ltm_column = quarterly_column + 1

        col_letter = get_column_letter(quarterly_column)
        ltm_letter = get_column_letter(ltm_column)
        logger.info(
            f"Q{quarter} {year}: Extracting values from column {col_letter} "
            f"(LTM {ltm_letter})"
        )
        
        values = {}
        
        # Helper function to extract value from cell, handling formulas
        def get_cell_value(sheet, row, col, wb_formulas=None, wb_data=None):
            """Get cell value, handling formulas by parsing and evaluating."""
            val = sheet.cell(row=row, column=col).value
            
            # If value is None and we have formulas workbook, try to parse formula
            if val is None and wb_formulas:
                formula_sheet = None
                for name in wb_formulas.sheetnames:
                    if 'Management Cijfers' in name:
                        formula_sheet = wb_formulas[name]
                        break
                
                if formula_sheet:
                    formula = formula_sheet.cell(row=row, column=col).value
                    if formula and isinstance(formula, str) and formula.startswith('='):
                        import re
                        from openpyxl.utils import column_index_from_string
                        
                        # Parse formula like: =-'BDO - Q3-25'!G76
                        match = re.match(r"=-?'([^']+)'!([A-Z]+\d+)", formula)
                        if match:
                            ref_sheet_name = match.group(1)
                            ref_cell = match.group(2)
                            
                            # Find referenced sheet in data workbook
                            if wb_data:
                                for sheet_name in wb_data.sheetnames:
                                    if ref_sheet_name in sheet_name:
                                        ref_sheet = wb_data[sheet_name]
                                        # Parse cell reference (e.g., G76)
                                        col_letter = ''.join(filter(str.isalpha, ref_cell))
                                        row_num = int(''.join(filter(str.isdigit, ref_cell)))
                                        col_num = column_index_from_string(col_letter)
                                        
                                        # Get value from referenced cell
                                        ref_val = ref_sheet.cell(row=row_num, column=col_num).value
                                        if ref_val is not None:
                                            # Apply negation (formula starts with -)
                                            return -ref_val
                        # Pattern: =SUM(AA23:AA24)
                        elif 'SUM' in formula:
                            match = re.search(r'SUM\(([A-Z]+)(\d+):([A-Z]+)(\d+)\)', formula)
                            if match:
                                start_col = column_index_from_string(match.group(1))
                                start_row = int(match.group(2))
                                end_col = column_index_from_string(match.group(3))
                                end_row = int(match.group(4))
                                
                                # Sum the range - recursively get values for each cell
                                total = 0
                                for r in range(start_row, end_row + 1):
                                    for c in range(start_col, end_col + 1):
                                        # Recursively get cell value (handles formulas)
                                        cell_val = get_cell_value(sheet, r, c, wb_formulas, wb_data)
                                        if cell_val is not None:
                                            total += cell_val
                                return total
            
            return val
        
        # Load formulas workbook for parsing if needed
        wb_formulas = None
        try:
            wb_formulas = openpyxl.load_workbook(ma_path, data_only=False)
        except:
            pass
        
        # Extract ALL values from Management Cijfers using the dynamically calculated column
        # Row 23: GTRI (stored as negative in formulas, take absolute)
        gtri_val = get_cell_value(cijfers_sheet, 23, quarterly_column, wb_formulas, wb)
        values['gtri'] = abs(gtri_val) / 1000 if gtri_val else 0
        
        # Row 24: Financial vacancy amount (stored as positive)
        vac_val = get_cell_value(cijfers_sheet, 24, quarterly_column, wb_formulas, wb)
        values['vacancy_amount'] = abs(vac_val) / 1000 if vac_val else 0
        
        # Row 25: Gross rental income (calculated value)
        gross_rental_val = get_cell_value(cijfers_sheet, 25, quarterly_column, wb_formulas, wb)
        values['gross_rental'] = abs(gross_rental_val) / 1000 if gross_rental_val else 0
        
        # Row 26: Vacancy percentage (calculated value, may be decimal like 0.054 or percentage like 5.4)
        vac_pct_val = get_cell_value(cijfers_sheet, 26, quarterly_column, wb_formulas, wb)
        if vac_pct_val is not None:
            # If value is < 1, it's a decimal (0.054), multiply by 100
            # If value is >= 1, it's already a percentage (5.4)
            if abs(vac_pct_val) < 1:
                values['vacancy_pct'] = abs(vac_pct_val) * 100
            else:
                values['vacancy_pct'] = abs(vac_pct_val)
        else:
            # Calculate from amount if percentage not available
            if values.get('gtri', 0) > 0:
                values['vacancy_pct'] = (values.get('vacancy_amount', 0) / values['gtri']) * 100
            else:
                values['vacancy_pct'] = 0
        
        # Row 32: Maintenance & repair costs
        maint_val = get_cell_value(cijfers_sheet, 32, quarterly_column, wb_formulas, wb)
        values['maintenance'] = abs(maint_val) / 1000 if maint_val else 0
        
        # Row 50: Cash proceeds sale (unit sale proceeds)
        cash_proceeds = get_cell_value(cijfers_sheet, 50, quarterly_column, wb_formulas, wb)
        values['unit_sales_proceeds'] = abs(cash_proceeds) / 1000 if cash_proceeds else 0
        
        # Extract LTM values from LTM column (quarterly_column + 1)
        gtri_ltm_val = get_cell_value(cijfers_sheet, 23, ltm_column, wb_formulas, wb)
        values['gtri_ltm'] = abs(gtri_ltm_val) / 1000 if gtri_ltm_val else 0
        
        vac_ltm_val = get_cell_value(cijfers_sheet, 24, ltm_column, wb_formulas, wb)
        values['vacancy_amount_ltm'] = abs(vac_ltm_val) / 1000 if vac_ltm_val else 0
        
        gross_rental_ltm_val = get_cell_value(cijfers_sheet, 25, ltm_column, wb_formulas, wb)
        values['gross_rental_ltm'] = abs(gross_rental_ltm_val) / 1000 if gross_rental_ltm_val else 0
        
        maint_ltm_val = get_cell_value(cijfers_sheet, 32, ltm_column, wb_formulas, wb)
        values['maintenance_ltm'] = abs(maint_ltm_val) / 1000 if maint_ltm_val else 0
        
        # Close formulas workbook if opened
        if wb_formulas:
            wb_formulas.close()
        
        # Fallback: If values are still None/zero (formulas not calculated), try BDO sheet directly
        if values['gtri'] == 0 or values['vacancy_amount'] == 0:
            logger.warning("Some values from Management Cijfers are None, trying BDO sheet as fallback")
            expected_bdo_name = f"BDO - Q{quarter}-{str(year)[-2:]}"
            for name in wb.sheetnames:
                if expected_bdo_name in name:
                    bdo_sheet = wb[name]
                    logger.info(f"Using BDO sheet as fallback: {name}")
                    
                    if values['gtri'] == 0:
                        gtri_q = bdo_sheet.cell(row=76, column=7).value
                        values['gtri'] = abs(gtri_q) / 1000 if gtri_q else 0
                        gtri_ltm = bdo_sheet.cell(row=76, column=8).value
                        values['gtri_ltm'] = abs(gtri_ltm) / 1000 if gtri_ltm else 0
                    
                    if values['vacancy_amount'] == 0:
                        vac_q = bdo_sheet.cell(row=77, column=7).value
                        values['vacancy_amount'] = abs(vac_q) / 1000 if vac_q else 0
                        vac_ltm = bdo_sheet.cell(row=77, column=8).value
                        values['vacancy_amount_ltm'] = abs(vac_ltm) / 1000 if vac_ltm else 0
                    
                    if values['maintenance'] == 0:
                        maint_q = bdo_sheet.cell(row=88, column=7).value
                        values['maintenance'] = abs(maint_q) / 1000 if maint_q else 0
                        maint_ltm = bdo_sheet.cell(row=88, column=8).value
                        values['maintenance_ltm'] = abs(maint_ltm) / 1000 if maint_ltm else 0
                    
                    # Recalculate derived values
                    values['gross_rental'] = values.get('gtri', 0) - values.get('vacancy_amount', 0)
                    values['gross_rental_ltm'] = values.get('gtri_ltm', 0) - values.get('vacancy_amount_ltm', 0)
                    if values.get('gtri', 0) > 0:
                        values['vacancy_pct'] = (values.get('vacancy_amount', 0) / values['gtri']) * 100
                    break
        
        # Log all extracted values
        logger.info(f"=== All Values from Management Cijfers Column {col_letter} (Q{quarter} {year}) ===")
        logger.info(f"  Gross Theoretical rental income: €{values['gtri']:,.1f}k")
        logger.info(f"  Gross rental income: €{values['gross_rental']:,.1f}k")
        logger.info(f"  Financial vacancy: {values['vacancy_pct']:.1f}% (€{values['vacancy_amount']:,.1f}k)")
        logger.info(f"  Maintenance: €{values['maintenance']:,.1f}k")
        logger.info(f"  Unit sale proceeds: €{values['unit_sales_proceeds']:,.0f}k")
        logger.info(f"  LTM GTRI: €{values['gtri_ltm']:,.1f}k")
        logger.info(f"  LTM Gross rental: €{values['gross_rental_ltm']:,.1f}k")
        logger.info(f"  LTM Maintenance: €{values['maintenance_ltm']:,.1f}k")
        
        wb.close()
        return values
        
    except Exception as e:
        logger.error(f"Error extracting values from Management Accounts: {e}")
        import traceback
        traceback.print_exc()
        return {}


_DUTCH_MONTHS = {
    1: 'januari', 2: 'februari', 3: 'maart', 4: 'april',
    5: 'mei', 6: 'juni', 7: 'juli', 8: 'augustus',
    9: 'september', 10: 'oktober', 11: 'november', 12: 'december',
}
_DUTCH_MONTHS_ALT = '|'.join(_DUTCH_MONTHS.values())
_ENGLISH_MONTHS = (
    'January|February|March|April|May|June|July|August|'
    'September|October|November|December'
)
_QUARTER_CANONICAL_RE = re.compile(r'Q([1-4])\s+(\d{4})')
_QUARTER_VARIANT_RES = [
    re.compile(r'Q([1-4])\s+(\d{4})'),
    re.compile(r'Q([1-4])[-/](\d{4})'),
    re.compile(r'Q([1-4])\s+(\d{2})\b'),
]
_ORDINAL_MAP = {
    '1': ('first', 'eerste'),
    '2': ('second', 'tweede'),
    '3': ('third', 'derde'),
    '4': ('fourth', 'vierde'),
}
_QUARTER_END_MONTHS = {
    '1': ('March', 'maart', '31', '3'),
    '2': ('June', 'juni', '30', '6'),
    '3': ('September', 'september', '30', '9'),
    '4': ('December', 'december', '31', '12'),
}


def format_report_date_dutch(when: datetime) -> str:
    """Format a datetime as a Dutch cover-page date (e.g. ``7 juni 2026``)."""
    return f"{when.day} {_DUTCH_MONTHS[when.month]} {when.year}"


def _strip_xml_to_text(xml_content: str) -> str:
    return re.sub(r'<[^>]+>', '', xml_content)


def read_docx_xml_text(docx_path: Path) -> str:
    """Concatenate visible text from document body, headers, and footers."""
    import zipfile

    parts: List[str] = []
    with zipfile.ZipFile(docx_path, 'r') as zin:
        for name in zin.namelist():
            if not name.endswith('.xml'):
                continue
            if not (
                name == 'word/document.xml'
                or '/header' in name
                or '/footer' in name
            ):
                continue
            parts.append(zin.read(name).decode('utf-8', errors='replace'))
    return _strip_xml_to_text(''.join(parts))


def discover_quarter_references(text: str, exclude: Optional[str] = None) -> List[str]:
    """Find unique quarter strings (``Q{n} YYYY``) anywhere in plain text."""
    found: set = set()
    for pattern in _QUARTER_VARIANT_RES:
        for match in pattern.finditer(text):
            q_num, year_part = match.group(1), match.group(2)
            year = year_part if len(year_part) == 4 else f"20{year_part}"
            canonical = f"Q{q_num} {year}"
            if exclude and canonical == exclude:
                continue
            found.add(canonical)
    return sorted(found, key=len, reverse=True)


def discover_quarter_references_in_docx(
    docx_path: Path, exclude: Optional[str] = None
) -> List[str]:
    return discover_quarter_references(read_docx_xml_text(docx_path), exclude=exclude)


def _parse_euro_k(text: str) -> float:
    return float(text.replace(',', '').replace(' ', ''))


def extract_prev_values_from_template(docx_path: Path) -> Dict[str, float]:
    """
    Read KPI values currently embedded in the Word template so replacements
    work regardless of how many reporting cycles the template has been through.
    """
    text = read_docx_xml_text(docx_path)
    prev: Dict[str, float] = {}

    patterns = [
        ('prev_gtri', r'GTRI[^€%]*€([\d,\.]+)\s*k'),
        ('prev_gross_rental_income', r'[Gg]ross rental income[^€%]*€([\d,\.]+)\s*k'),
        ('prev_vacancy_pct', r'[Ff]inancial vacancy[^%]*?([\d,\.]+)\s*%'),
        ('prev_vacancy_amount', r'[Ff]inancial vacancy[^€]*€([\d,\.]+)\s*k'),
        ('prev_gtri_ltm', r'(?:rent roll|yields)[^€]*€([\d,\.]+)\s*k'),
        ('prev_maintenance', r'[Mm]aintenance[^€]*€([\d,\.]+)\s*k'),
        ('prev_unit_sales_proceeds', r'(?:Unit sale proceeds|sale proceeds)[^€]*€([\d,\.]+)\s*k'),
    ]
    for key, pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if not match:
            continue
        raw = match.group(1)
        if key == 'prev_vacancy_pct':
            prev[key] = float(raw.replace(',', '.'))
        else:
            prev[key] = _parse_euro_k(raw)

    return prev


def build_quarter_replacement_pairs(
    stale_quarters: List[str], current_quarter: str
) -> List[Tuple[str, str]]:
    """Build (old, new) replacement pairs for every stale quarter string."""
    pairs: List[Tuple[str, str]] = []
    seen: set = set()

    def add(old: str, new: str) -> None:
        if not old or old == new or (old, new) in seen:
            return
        seen.add((old, new))
        pairs.append((old, new))

    for stale in stale_quarters:
        if stale == current_quarter:
            continue
        stale_parts = stale.split()
        current_parts = current_quarter.split()
        if len(stale_parts) != 2 or len(current_parts) != 2:
            add(stale, current_quarter)
            continue

        old_q_num = stale_parts[0][1]
        old_year = stale_parts[1]
        new_q_num = current_parts[0][1]
        new_year = current_parts[1]

        add(stale, current_quarter)
        add(stale.replace(' ', '-'), current_quarter.replace(' ', '-'))
        add(stale.replace(' ', '/'), current_quarter.replace(' ', '/'))
        add(f"Q{old_q_num} {old_year[-2:]}", f"Q{new_q_num} {new_year[-2:]}")
        add(f"Q{old_q_num}-{old_year[-2:]}", f"Q{new_q_num}-{new_year[-2:]}")
        add(f"{old_year[-2:]}Q{old_q_num}", f"{new_year[-2:]}Q{new_q_num}")
        add(f"{old_year} – Q{old_q_num}", f"{new_year} – Q{new_q_num}")
        add(f"{old_year} - Q{old_q_num}", f"{new_year} - Q{new_q_num}")

        if old_q_num in _ORDINAL_MAP and new_q_num in _ORDINAL_MAP:
            old_en, old_nl = _ORDINAL_MAP[old_q_num]
            new_en, new_nl = _ORDINAL_MAP[new_q_num]
            add(f"{old_en} quarter", f"{new_en} quarter")
            add(f"{old_en.capitalize()} quarter", f"{new_en.capitalize()} quarter")
            add(f"{old_nl} kwartaal", f"{new_nl} kwartaal")

        if old_q_num in _QUARTER_END_MONTHS and new_q_num in _QUARTER_END_MONTHS:
            old_month, old_month_nl, old_day, old_month_num = _QUARTER_END_MONTHS[old_q_num]
            new_month, new_month_nl, new_day, new_month_num = _QUARTER_END_MONTHS[new_q_num]
            add(old_month, new_month)
            add(old_month.lower(), new_month.lower())
            add(old_month_nl, new_month_nl)
            add(
                f"{old_day} {old_month} {old_year}",
                f"{new_day} {new_month} {new_year}",
            )
            add(
                f"{old_day} {old_month_nl} {old_year}",
                f"{new_day} {new_month_nl} {new_year}",
            )
            add(
                f"{old_day}-{old_month_num}-{old_year}",
                f"{new_day}-{new_month_num}-{new_year}",
            )
            add(
                f"1-{int(old_month_num) + 1}-{old_year}",
                f"1-{int(new_month_num) + 1}-{new_year}",
            )

    pairs.sort(key=lambda item: len(item[0]), reverse=True)
    return pairs


def discover_cover_date_strings(docx_path: Path, report_date: str) -> List[str]:
    """Find cover-page date strings in the template that should become ``report_date``."""
    text = read_docx_xml_text(docx_path)
    dates: set = set()
    for pattern in (
        rf'\d{{1,2}}\s+(?:{_DUTCH_MONTHS_ALT})\s+\d{{4}}',
        rf'\d{{1,2}}\s+(?:{_ENGLISH_MONTHS})\s+\d{{4}}',
    ):
        for match in re.finditer(pattern, text, re.IGNORECASE):
            candidate = match.group(0)
            if candidate.lower() != report_date.lower():
                dates.add(candidate)
    return sorted(dates, key=len, reverse=True)


def build_report_values(
    *,
    ma_values: Dict[str, float],
    rent_roll_k: float,
    rent_roll_units: int,
    units_sold_quarter: int,
    unit_sales_proceeds: float,
    word_template_path: Path,
    report_date: Optional[datetime] = None,
) -> ReportValues:
    """Assemble ``ReportValues`` with template-derived previous KPI anchors."""
    prev = extract_prev_values_from_template(word_template_path)
    when = report_date or datetime.now()
    gtri_ltm = ma_values.get('gtri_ltm') or 0.0
    if gtri_ltm <= 0:
        gtri_ltm = rent_roll_k

    return ReportValues(
        report_date=format_report_date_dutch(when),
        gtri=ma_values.get('gtri', 0.0),
        gtri_ltm=gtri_ltm,
        gross_rental_income=ma_values.get('gross_rental', 0.0),
        gross_rental_income_ltm=ma_values.get('gross_rental_ltm', 0.0),
        financial_vacancy_pct=ma_values.get('vacancy_pct', 0.0),
        financial_vacancy_amount=ma_values.get('vacancy_amount', 0.0),
        rent_roll_annual=rent_roll_k,
        rent_roll_units=rent_roll_units,
        units_sold_quarter=units_sold_quarter,
        unit_sales_proceeds=unit_sales_proceeds,
        maintenance_amount=ma_values.get('maintenance', 0.0),
        maintenance_ltm=ma_values.get('maintenance_ltm', 0.0),
        capex_amount=0.0,
        capex_ltm=0.0,
        unit_sales_narrative=(
            f"{units_sold_quarter} unit(s) sold for €{unit_sales_proceeds:,.0f}k"
            if units_sold_quarter > 0
            else "No unit sales this quarter."
        ),
        maintenance_detail="",
        sustainability_detail="",
        prev_gtri=prev.get('prev_gtri', 0.0),
        prev_gtri_ltm=prev.get('prev_gtri_ltm', 0.0),
        prev_gross_rental_income=prev.get('prev_gross_rental_income', 0.0),
        prev_vacancy_pct=prev.get('prev_vacancy_pct', 0.0),
        prev_vacancy_amount=prev.get('prev_vacancy_amount', 0.0),
        prev_rent_roll=0.0,
        prev_maintenance=prev.get('prev_maintenance', 0.0),
        prev_unit_sales_proceeds=prev.get('prev_unit_sales_proceeds', 0.0),
    )


def _numeric_variants(value: float) -> List[float]:
    if value <= 0:
        return []
    variants = {round(value, 1), round(value, 0), value}
    if value >= 100:
        variants.add(round(value, -1))
    return [v for v in variants if v > 0]


def _collapse_xml_runs_to_single_value(xml_span: str, new_text: str) -> str:
    first = True

    def repl(_match: re.Match) -> str:
        nonlocal first
        if first:
            first = False
            return f'>{new_text}</w:t>'
        return '></w:t>'

    return re.sub(r'>([^<]*)</w:t>', repl, xml_span)


def _find_fragmented_number_span(content: str, target_value: float, max_runs: int = 12) -> Optional[str]:
    """Locate an XML span whose concatenated ``w:t`` text equals ``target_value``."""
    if target_value <= 0:
        return None

    runs = list(re.finditer(r'>([^<]*)</w:t>', content))
    targets = _numeric_variants(target_value)

    for i in range(len(runs)):
        combined = ''
        for j in range(i, min(i + max_runs, len(runs))):
            combined += runs[j].group(1)
            normalized = (
                combined.replace('€', '')
                .replace(' ', '')
                .replace(',', '')
                .replace('k', '')
                .replace('m', '')
                .replace('%', '')
            )
            if not normalized or normalized in ('.', '-'):
                continue
            try:
                parsed = float(normalized)
            except ValueError:
                continue
            for candidate in targets:
                if abs(parsed - candidate) <= 0.25:
                    return content[runs[i].start():runs[j].end()]
    return None


@dataclass 
class NumericValueMapping:
    """Maps a KPI to its context pattern and value formatter."""
    kpi_name: str
    # Regex pattern to find the value in context (must have a capture group for the number)
    context_pattern: str
    # How to format the new value
    formatter: str  # 'euro_k', 'euro_m', 'percent', 'integer'
    # Which ReportValues attribute to use
    value_attr: str
    # Optional: previous value attribute for better matching
    prev_value_attr: Optional[str] = None


class WordTemplateUpdater:
    """
    Updates Word template for quarterly report.
    
    Uses pandoc for conversion to/from markdown for text operations,
    and python-docx for structural operations.
    """
    
    # Placeholder patterns in the template
    PLACEHOLDERS = {
        '{{GTRI}}': 'gtri',
        '{{GROSS_RENTAL_INCOME}}': 'gross_rental_income',
        '{{RENT_ROLL}}': 'rent_roll_annual',
        '{{VACANCY}}': 'financial_vacancy_pct',
        '{{UNITS_SOLD}}': 'units_sold_quarter',
        '{{MAINTENANCE}}': 'maintenance_amount',
        '{{CAPEX}}': 'capex_amount',
    }
    
    def __init__(self, template_path: str, output_path: str):
        self.template_path = Path(template_path)
        self.output_path = Path(output_path)
        # Use system temp directory
        self.working_dir = Path(tempfile.gettempdir()) / "word_update"
        self.working_dir.mkdir(parents=True, exist_ok=True)
        
    def update(self, 
               values: ReportValues, 
               previous_quarter: str,
               current_quarter: str) -> Path:
        """
        Update template with new values using pandoc.
        Simpler method but may lose some formatting.
        
        Args:
            values: ReportValues with all data to insert
            previous_quarter: Previous quarter string (e.g., "Q2 2025")
            current_quarter: Current quarter string (e.g., "Q3 2025")
            
        Returns:
            Path to updated Word document
        """
        logger.info(f"Updating Word template for {current_quarter} (pandoc method)")
        
        # Step 1: Convert to markdown for text operations
        md_path = self.working_dir / "template.md"
        self._docx_to_markdown(self.template_path, md_path)
        
        # Step 2: Read markdown content
        with open(md_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Step 3: Global quarter replacement
        content = self._replace_quarter(content, previous_quarter, current_quarter)
        
        # Step 4: Replace placeholders with values
        content = self._replace_placeholders(content, values)
        
        # Step 5: Write updated markdown
        updated_md = self.working_dir / "updated.md"
        with open(updated_md, 'w', encoding='utf-8') as f:
            f.write(content)
        
        # Step 6: Convert back to docx (preserving reference doc styling)
        self._markdown_to_docx(updated_md, self.output_path, self.template_path)
        
        logger.info(f"Generated updated report: {self.output_path}")
        return self.output_path
    
    def _docx_to_markdown(self, docx_path: Path, md_path: Path):
        """Convert Word document to markdown using pandoc."""
        cmd = [
            'pandoc',
            str(docx_path),
            '-o', str(md_path),
            '--wrap=none'  # Preserve line structure
        ]
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode != 0:
            raise RuntimeError(f"Pandoc conversion failed: {result.stderr}")
    
    def _markdown_to_docx(self, md_path: Path, docx_path: Path, reference_doc: Path):
        """Convert markdown back to Word using pandoc with reference doc."""
        cmd = [
            'pandoc',
            str(md_path),
            '-o', str(docx_path),
            f'--reference-doc={reference_doc}'  # Use original as style reference
        ]
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode != 0:
            raise RuntimeError(f"Pandoc conversion failed: {result.stderr}")
    
    def _replace_quarter(self, content: str, old_quarter: str, new_quarter: str) -> str:
        """
        Replace ALL quarter references comprehensively.
        
        This addresses Issue 7.1 from feedback: mixed quarter references.
        Handles all possible formats and variations.
        """
        import re
        
        # Extract quarter number and year
        old_parts = old_quarter.split()  # e.g., ["Q2", "2025"]
        new_parts = new_quarter.split()
        
        if len(old_parts) != 2 or len(new_parts) != 2:
            logger.warning(f"Unexpected quarter format: {old_quarter} -> {new_quarter}")
            return content
        
        old_q_num = old_parts[0][1]  # "2" from "Q2"
        old_year = old_parts[1]  # "2025"
        new_q_num = new_parts[0][1]  # "3" from "Q3"
        new_year = new_parts[1]  # "2025"
        
        # Handle ordinal words (second quarter → third quarter, etc.)
        ordinal_map = {
            '1': ('first', 'eerste'),
            '2': ('second', 'tweede'),
            '3': ('third', 'derde'),
            '4': ('fourth', 'vierde'),
        }
        if old_q_num in ordinal_map and new_q_num in ordinal_map:
            old_ordinal_en, old_ordinal_nl = ordinal_map[old_q_num]
            new_ordinal_en, new_ordinal_nl = ordinal_map[new_q_num]
            
            # English ordinal patterns - case variations
            content = re.sub(rf'\b{old_ordinal_en}\s+quarter\b', f'{new_ordinal_en} quarter', content, flags=re.IGNORECASE)
            content = re.sub(rf'\bthe\s+{old_ordinal_en}\b', f'the {new_ordinal_en}', content, flags=re.IGNORECASE)
            
            # Dutch ordinal patterns
            content = re.sub(rf'\b{old_ordinal_nl}\s+kwartaal\b', f'{new_ordinal_nl} kwartaal', content, flags=re.IGNORECASE)
        
        # Comprehensive replacement patterns
        replacements = [
            # Standard formats
            (old_quarter, new_quarter),  # Q2 2025 -> Q3 2025
            (old_quarter.lower(), new_quarter.lower()),  # q2 2025 -> q3 2025
            (old_quarter.replace(' ', '-'), new_quarter.replace(' ', '-')),  # Q2-2025
            (old_quarter.replace(' ', '/'), new_quarter.replace(' ', '/')),  # Q2/2025
            (f"Q{old_q_num} {old_year[-2:]}", f"Q{new_q_num} {new_year[-2:]}"),  # Q2 25
            (f"Q{old_q_num}-{old_year[-2:]}", f"Q{new_q_num}-{new_year[-2:]}"),  # Q2-25
            (f"{old_year[-2:]}Q{old_q_num}", f"{new_year[-2:]}Q{new_q_num}"),  # 25Q2
            
            # Year-quarter formats with different separators
            (f"{old_year} – Q{old_q_num}", f"{new_year} – Q{new_q_num}"),  # 2025 – Q2
            (f"{old_year} - Q{old_q_num}", f"{new_year} - Q{new_q_num}"),  # 2025 - Q2
            
            # Title and header formats (as seen in Sample.pdf)
            (f"Quarterly Report {old_quarter}", f"Quarterly Report {new_quarter}"),
            (f"Quarterly QSP - {old_quarter}", f"Quarterly QSP - {new_quarter}"),
            (f"in {old_quarter}", f"in {new_quarter}"),
            
            # Text variations
            (f"Actions {old_quarter}", f"Actions {new_quarter}"),
            (f"Portfolio highlights – {old_quarter}", f"Portfolio highlights – {new_quarter}"),
            (f"Portfolio highlights - {old_quarter}", f"Portfolio highlights - {new_quarter}"),
            (f"Report {old_quarter}", f"Report {new_quarter}"),
            (f"Quarter {old_q_num} {old_year}", f"Quarter {new_q_num} {new_year}"),
            
            # Compliance certificate references
            (f"Interest Period: {self._get_period_date(old_q_num, old_year)}", 
             f"Interest Period: {self._get_period_date(new_q_num, new_year)}"),
            
            # Dutch variations
            (f"Kwartaal {old_q_num} {old_year}", f"Kwartaal {new_q_num} {new_year}"),
        ]
        
        for old_pattern, new_pattern in replacements:
            content = content.replace(old_pattern, new_pattern)
            # Also try case-insensitive for text variations
            if not old_pattern[0].isdigit():
                content = re.sub(re.escape(old_pattern), new_pattern, content, flags=re.IGNORECASE)
        
        # Handle month references
        quarter_to_month = {
            'Q1': ('March', '31', 'maart'),
            'Q2': ('June', '30', 'juni'),
            'Q3': ('September', '30', 'september'),
            'Q4': ('December', '31', 'december'),
        }
        
        old_q = old_parts[0]
        new_q = new_parts[0]
        
        if old_q in quarter_to_month and new_q in quarter_to_month:
            old_month, old_day, old_month_nl = quarter_to_month[old_q]
            new_month, new_day, new_month_nl = quarter_to_month[new_q]
            
            # Replace English months
            content = content.replace(old_month, new_month)
            content = content.replace(old_month.lower(), new_month.lower())
            
            # Replace Dutch months
            content = content.replace(old_month_nl, new_month_nl)
            content = content.replace(old_month_nl.capitalize(), new_month_nl.capitalize())
            
            # Replace date formats like "30 June 2025" and "30-6-2025"
            for year in [old_year, str(int(old_year) - 1), str(int(old_year) + 1)]:
                # English format: 30 June 2025
                old_date = f"{old_day} {old_month} {year}"
                new_date = f"{new_day} {new_month} {new_year}"
                content = content.replace(old_date, new_date)
                
                # Dutch format: 30 juni 2025
                old_date_nl = f"{old_day} {old_month_nl} {year}"
                new_date_nl = f"{new_day} {new_month_nl} {new_year}"
                content = content.replace(old_date_nl, new_date_nl)
                
                # Numeric formats: 30-6-2025, 30/6/2025
                old_month_num = {'March': 3, 'June': 6, 'September': 9, 'December': 12}[old_month]
                new_month_num = {'March': 3, 'June': 6, 'September': 9, 'December': 12}[new_month]
                
                content = content.replace(f"{old_day}-{old_month_num}-{year}", f"{new_day}-{new_month_num}-{new_year}")
                content = content.replace(f"{old_day}/{old_month_num}/{year}", f"{new_day}/{new_month_num}/{new_year}")
                
                # Rent roll date format: 1-10-2025 (1st of next month after quarter)
                next_month_old = old_month_num + 1 if old_month_num < 12 else 1
                next_year_old = year if old_month_num < 12 else str(int(year) + 1)
                next_month_new = new_month_num + 1 if new_month_num < 12 else 1
                next_year_new = new_year if new_month_num < 12 else str(int(new_year) + 1)
                
                content = content.replace(f"1-{next_month_old}-{next_year_old}", f"1-{next_month_new}-{next_year_new}")
        
        return content
    
    def _get_period_date(self, q_num: str, year: str) -> str:
        """Get period end date string for a quarter."""
        month_days = {'1': ('3', '31'), '2': ('6', '30'), '3': ('9', '30'), '4': ('12', '31')}
        if q_num in month_days:
            month, day = month_days[q_num]
            return f"{day}-{month}-{year}"
        return ""
    
    def _replace_placeholders(self, content: str, values: ReportValues) -> str:
        """Replace placeholder values in content."""
        replacements = {
            '{{GTRI}}': f"€{values.gtri:,.1f}k",
            '{{GROSS_RENTAL_INCOME}}': f"€{values.gross_rental_income:,.1f}k",
            '{{RENT_ROLL}}': f"€{values.rent_roll_annual:,.1f}k",
            '{{VACANCY}}': f"{values.financial_vacancy_pct:.1f}%",
            '{{UNITS_SOLD}}': str(values.units_sold_quarter),
            '{{MAINTENANCE}}': f"€{values.maintenance_amount:,.0f}k",
            '{{CAPEX}}': f"€{values.capex_amount:,.0f}k",
            '{{REPORT_DATE}}': values.report_date,
        }
        
        for placeholder, value in replacements.items():
            content = content.replace(placeholder, value)
        
        # Handle narrative sections
        if values.unit_sales_narrative:
            content = content.replace('{{UNIT_SALES_NARRATIVE}}', values.unit_sales_narrative)
        
        if values.maintenance_detail:
            content = content.replace('{{MAINTENANCE_DETAIL}}', values.maintenance_detail)
        
        if values.sustainability_detail:
            content = content.replace('{{SUSTAINABILITY_DETAIL}}', values.sustainability_detail)
        
        return content
    
    def update_with_python_docx(self, values: ReportValues, 
                                 previous_quarter: str, 
                                 current_quarter: str) -> Path:
        """
        Update Word document including text boxes using direct XML manipulation.
        
        This method handles:
        - Regular paragraphs (via python-docx)
        - Text boxes and shapes (via direct XML manipulation)
        - Headers and footers
        
        Args:
            values: ReportValues with all data to insert
            previous_quarter: Previous quarter string (e.g., "Q2 2025")
            current_quarter: Current quarter string (e.g., "Q3 2025")
            
        Returns:
            Path to updated Word document
        """
        from docx import Document
        import zipfile
        import os
        
        logger.info(f"Updating Word template for {current_quarter} (python-docx + XML method)")
        
        stale_quarters = discover_quarter_references_in_docx(
            self.template_path, exclude=current_quarter
        )
        if previous_quarter not in stale_quarters and previous_quarter != current_quarter:
            stale_quarters.append(previous_quarter)
        self._quarter_pairs = build_quarter_replacement_pairs(
            stale_quarters, current_quarter
        )
        self._cover_date_patterns = discover_cover_date_strings(
            self.template_path, values.report_date
        )
        self._cover_date_patterns.extend(self._get_old_date_patterns(previous_quarter))
        seen_dates = set()
        self._cover_date_patterns = [
            p for p in self._cover_date_patterns
            if p.lower() != values.report_date.lower()
            and not (p.lower() in seen_dates or seen_dates.add(p.lower()))
        ]
        self._cover_date_patterns.sort(key=len, reverse=True)
        
        # Copy template to output location
        shutil.copy2(self.template_path, self.output_path)
        
        # Step 1: Update regular paragraphs/tables via python-docx (basic text replacements)
        doc = Document(self.output_path)
        
        # Process paragraphs
        for para in doc.paragraphs:
            self._process_paragraph(para, values, previous_quarter, current_quarter)
            self._update_cover_page_date(para, self._cover_date_patterns, values.report_date)
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._process_paragraph(para, values, previous_quarter, current_quarter)
                        self._update_cover_page_date(para, self._cover_date_patterns, values.report_date)
        
        # Process headers/footers
        for section in doc.sections:
            for para in section.header.paragraphs:
                self._process_paragraph(para, values, previous_quarter, current_quarter)
                self._update_cover_page_date(para, self._cover_date_patterns, values.report_date)
            for para in section.footer.paragraphs:
                self._process_paragraph(para, values, previous_quarter, current_quarter)
                self._update_cover_page_date(para, self._cover_date_patterns, values.report_date)
        
        doc.save(self.output_path)
        
        # Step 2: After python-docx saves, apply direct XML manipulation for:
        # - Text boxes (which python-docx can't access)
        # - Fragmented text replacements
        # - Context-aware numeric replacements (€ values split across multiple XML elements)
        # This must happen AFTER python-docx saves, otherwise doc.save() overwrites our changes
        text_replacements = self._build_replacement_map(values, previous_quarter, current_quarter)
        self._update_text_boxes_xml(text_replacements, values)
        
        logger.info(f"Generated updated report: {self.output_path}")
        return self.output_path
    
    def _build_replacement_map(self, values: ReportValues, old_quarter: str, new_quarter: str) -> List[Tuple[str, str]]:
        """
        Build comprehensive replacement map for all text that needs updating.
        
        Returns list of (old_text, new_text) tuples, sorted longest-first.
        
        INCLUDES: 
        - Quarter text replacements (all stale quarters found in template)
        - Date replacements
        - NUMERIC VALUE REPLACEMENTS (Page 4 KPIs)
        """
        replacements: List[Tuple[str, str]] = []

        stale_quarters = discover_quarter_references_in_docx(
            self.template_path, exclude=new_quarter
        )
        if old_quarter not in stale_quarters and old_quarter != new_quarter:
            stale_quarters.append(old_quarter)
        replacements.extend(build_quarter_replacement_pairs(stale_quarters, new_quarter))

        numeric_replacements = self._build_numeric_replacements(values)
        replacements.extend(numeric_replacements)

        for old_date in discover_cover_date_strings(self.template_path, values.report_date):
            replacements.append((old_date, values.report_date))

        old_date_patterns = self._get_old_date_patterns(old_quarter)
        for old_date in old_date_patterns:
            if old_date.lower() != values.report_date.lower():
                replacements.append((old_date, values.report_date))

        replacements.sort(key=lambda x: len(x[0]), reverse=True)
        return replacements
    
    def _build_numeric_replacements(self, values: ReportValues) -> List[Tuple[str, str]]:
        """
        Build replacement tuples for numeric values in the Word document.
        
        The Word document contains financial KPIs embedded in narrative text.
        We need to find patterns like "€3,200k" or "5.4%" and replace them
        with the new quarter's values.
        
        This method generates replacement pairs for common numeric formats
        that might appear in the document.
        """
        replacements = []
        
        # Helper functions for formatting
        def format_euro_k(value: float, decimals: int = 0) -> List[str]:
            """Generate multiple format variations for Euro amounts in thousands."""
            formats = []
            if value == 0:
                return formats
            
            abs_val = abs(value)
            
            # Various formatting styles that might appear in the document
            if decimals == 0:
                formats.extend([
                    f"€{abs_val:,.0f}k",
                    f"€{abs_val:,.0f} k",
                    f"€ {abs_val:,.0f}k",
                    f"€{abs_val:,.0f}",  # Without k suffix
                    f"€{int(abs_val):,}k",
                    f"€{int(abs_val):,}",
                ])
            else:
                formats.extend([
                    f"€{abs_val:,.{decimals}f}k",
                    f"€{abs_val:,.{decimals}f} k", 
                    f"€ {abs_val:,.{decimals}f}k",
                    f"€{abs_val:,.{decimals}f}",
                ])
            
            # European format with dot as thousand separator
            if abs_val >= 1000:
                euro_fmt = f"{abs_val:,.0f}".replace(',', '.')
                formats.extend([
                    f"€{euro_fmt}k",
                    f"€{euro_fmt}",
                ])
            
            return formats
        
        def format_euro_m(value: float) -> List[str]:
            """Generate format variations for Euro amounts in millions."""
            formats = []
            if value == 0:
                return formats
            
            abs_val = abs(value)
            formats.extend([
                f"€{abs_val:.1f}m",
                f"€{abs_val:.1f} m",
                f"€ {abs_val:.1f}m",
                f"€{abs_val:.2f}m",
            ])
            return formats
        
        def format_percent(value: float) -> List[str]:
            """Generate format variations for percentages."""
            formats = []
            abs_val = abs(value)
            formats.extend([
                f"{abs_val:.1f}%",
                f"{abs_val:.1f} %",
                f"{abs_val:.2f}%",
                f"{int(abs_val)}%",
            ])
            return formats
        
        # === PAGE 4 NUMERIC REPLACEMENTS ===
        # These are the key KPIs that need to be updated
        
        # 1. GTRI (Gross Theoretical Rental Income) - quarterly value
        # Format: typically "€3,200k" or "€3.2m"
        if values.gtri > 0:
            gtri_k = values.gtri  # Already in thousands
            
            # If we have previous GTRI value, create direct replacement
            if values.prev_gtri > 0:
                for old_fmt in format_euro_k(values.prev_gtri, 0):
                    for new_fmt in format_euro_k(gtri_k, 0)[:1]:  # Just use first format
                        replacements.append((old_fmt, new_fmt))
        
        # 2. Gross Rental Income - quarterly value
        if values.gross_rental_income > 0 and values.prev_gross_rental_income > 0:
            for old_fmt in format_euro_k(values.prev_gross_rental_income, 0):
                for new_fmt in format_euro_k(values.gross_rental_income, 0)[:1]:
                    replacements.append((old_fmt, new_fmt))
        
        # 3. Financial Vacancy percentage
        if values.prev_vacancy_pct > 0:
            for old_fmt in format_percent(values.prev_vacancy_pct):
                for new_fmt in format_percent(values.financial_vacancy_pct)[:1]:
                    replacements.append((old_fmt, new_fmt))
        
        # 4. Rent Roll total
        if values.prev_rent_roll > 0 and values.rent_roll_annual > 0:
            for old_fmt in format_euro_k(values.prev_rent_roll, 0):
                for new_fmt in format_euro_k(values.rent_roll_annual, 0)[:1]:
                    replacements.append((old_fmt, new_fmt))
        
        # 5. Maintenance amount
        if values.prev_maintenance > 0 and values.maintenance_amount > 0:
            for old_fmt in format_euro_k(values.prev_maintenance, 0):
                for new_fmt in format_euro_k(values.maintenance_amount, 0)[:1]:
                    replacements.append((old_fmt, new_fmt))
        
        # 6. Unit sales proceeds
        if values.prev_unit_sales_proceeds > 0 and values.unit_sales_proceeds > 0:
            for old_fmt in format_euro_k(values.prev_unit_sales_proceeds, 0):
                for new_fmt in format_euro_k(values.unit_sales_proceeds, 0)[:1]:
                    replacements.append((old_fmt, new_fmt))
        
        logger.info(f"Built {len(replacements)} numeric value replacements")
        return replacements
    
    def _update_text_boxes_xml(self, replacements: List[Tuple[str, str]], values: ReportValues = None):
        """
        Update text in text boxes by directly manipulating the docx XML.
        
        Text boxes are stored as drawing elements in the document.xml,
        and python-docx can't access them directly.
        
        IMPORTANT: Word XML often splits text across multiple <w:t> elements,
        e.g., "Q2 2025" might be stored as ["Q2", " 2025"]. We handle this
        by using regex patterns that account for XML tags between text fragments.
        
        Also applies context-aware numeric value replacements for Page 4 KPIs.
        """
        import zipfile
        import os
        
        # Read the docx as a zip file
        temp_dir = self.working_dir / "docx_xml"
        temp_dir.mkdir(parents=True, exist_ok=True)
        
        # Extract all files
        with zipfile.ZipFile(self.output_path, 'r') as zin:
            zin.extractall(temp_dir)
        
        # Process document.xml and any other relevant XML files
        xml_files = [
            temp_dir / "word" / "document.xml",
        ]
        
        # Also process header/footer files if they exist
        word_dir = temp_dir / "word"
        for f in word_dir.iterdir():
            if f.suffix == '.xml' and ('header' in f.name or 'footer' in f.name):
                xml_files.append(f)
        
        updated_count = 0
        numeric_updates = 0
        for xml_file in xml_files:
            if xml_file.exists():
                with open(xml_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                original_content = content
                
                # Build fragmented patterns that account for XML tags between text pieces
                # E.g., "Q2 2025" might be: <w:t>Q2</w:t></w:r><w:r><w:t> 2025</w:t>
                content = self._apply_fragmented_replacements(content, replacements)
                
                # Apply numeric value replacements for Page 4 KPIs
                if values:
                    # Apply fragmented euro value replacements FIRST
                    # This handles Portfolio Highlights values (4 + . + 1 + % as separate elements)
                    # before _apply_context_numeric_replacements modifies body text
                    content, frag_count = self._apply_fragmented_euro_replacements(content, values)
                    numeric_updates += frag_count
                    
                    # Apply context-aware numeric replacements
                    # This handles body text values (4. + 1 as combined elements)
                    content, num_count = self._apply_context_numeric_replacements(content, values)
                    numeric_updates += num_count
                    
                    # Apply fragmented date replacements (e.g., "1- 7 -202 5" -> "1- 10 -202 5")
                    content = self._apply_fragmented_date_replacements(content, replacements)
                
                # Resize Portfolio Highlights box to ensure all content is visible
                content = self._resize_portfolio_highlights_box(content)
                
                if content != original_content:
                    with open(xml_file, 'w', encoding='utf-8') as f:
                        f.write(content)
                    updated_count += 1
        
        if numeric_updates > 0:
            logger.info(f"Applied {numeric_updates} context-aware numeric replacements")
        
        logger.info(f"Updated {updated_count} XML files in docx")
        
        # Repack the docx
        with zipfile.ZipFile(self.output_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for root, dirs, files in os.walk(temp_dir):
                for file in files:
                    file_path = Path(root) / file
                    arcname = file_path.relative_to(temp_dir)
                    zout.write(file_path, arcname)
        
        # Cleanup temp directory
        shutil.rmtree(temp_dir, ignore_errors=True)
    
    def _resize_portfolio_highlights_box(self, content: str) -> str:
        """
        Adjust the Portfolio Highlights box to fit all content.
        
        Instead of resizing the box (which can overlap other content), this method:
        1. Keeps the original box size
        2. Reduces the font size of text inside the box so all content fits
        
        Args:
            content: The XML content of document.xml
            
        Returns:
            Updated XML content with smaller font sizes in Portfolio Highlights box
        """
        import re
        
        # Find all anchors individually (non-greedy pattern that stops at first closing tag)
        anchor_pattern = r'<wp:anchor[^>]*>(?:(?!</wp:anchor>).)*?</wp:anchor>'
        
        def is_portfolio_highlights_box(anchor_xml: str) -> bool:
            """Check if this anchor is the Portfolio Highlights box."""
            return ('rental income:' in anchor_xml and 
                    ('3,273' in anchor_xml or '3,095' in anchor_xml or 'vacancy:' in anchor_xml))
        
        def reduce_font_size(anchor_xml: str) -> str:
            """Reduce font sizes in the Portfolio Highlights box."""
            updated = False
            
            # Find all font size declarations in the box and reduce them
            # Word uses w:sz (font size in half-points, so 20 = 10pt, 14 = 7pt)
            # Current sizes are around 14-20 (7-10pt), reduce to 12-16 (6-8pt)
            
            # Pattern to find w:sz elements
            sz_pattern = r'(<w:sz\s+w:val=")(\d+)(")'
            szCs_pattern = r'(<w:szCs\s+w:val=")(\d+)(")'
            
            def reduce_size(match):
                prefix = match.group(1)
                size = int(match.group(2))
                suffix = match.group(3)
                
                # Reduce font size by about 30% (but keep minimum of 10 = 5pt)
                new_size = max(10, int(size * 0.7))
                return f'{prefix}{new_size}{suffix}'
            
            # Apply font size reduction
            new_anchor_xml = re.sub(sz_pattern, reduce_size, anchor_xml)
            new_anchor_xml = re.sub(szCs_pattern, reduce_size, new_anchor_xml)
            
            if new_anchor_xml != anchor_xml:
                updated = True
                logger.info("Reduced font sizes in Portfolio Highlights box")
            
            return new_anchor_xml
        
        # Find all anchors and process only the Portfolio Highlights box
        updated_content = content
        anchors_found = 0
        
        for match in reversed(list(re.finditer(anchor_pattern, content, flags=re.DOTALL))):
            anchor_xml = match.group(0)
            
            if is_portfolio_highlights_box(anchor_xml):
                anchors_found += 1
                modified = reduce_font_size(anchor_xml)
                updated_content = updated_content[:match.start()] + modified + updated_content[match.end():]
        
        if anchors_found > 0:
            logger.info(f"Adjusted font sizes in {anchors_found} Portfolio Highlights box(es)")
        else:
            logger.warning("Could not find Portfolio Highlights box in document")
        
        return updated_content
    
    def _apply_context_numeric_replacements(self, content: str, values: ReportValues) -> Tuple[str, int]:
        """
        Apply context-aware numeric value replacements in the Word XML.
        
        This method finds numeric values (€ amounts, percentages) that appear near
        specific KPI context words (GTRI, rental income, vacancy, etc.) and replaces
        them with the new quarter's values.
        
        Word XML fragments numbers across multiple <w:t> elements, e.g.:
        "€3,200k" might be: <w:t>€</w:t><w:t>3,</w:t><w:t>200</w:t><w:t>k</w:t>
        
        The document has specific Q2 values that need to be replaced with Q3 values:
        - GTRI: €3,200.6k → €3,273.6k
        - Financial vacancy: €106.3k → €177.5k (amount), 4.1% → 5.4% (percentage)
        - Gross rental income: €3,067.5k → €3,095.9k
        - Rent roll yields: €12,940.2k → €13,317.9k (this is LTM GTRI)
        - Unit sales proceeds: €762.5k → €203.5k
        - Maintenance: €297k → €244k
        
        Returns:
            (updated_content, count_of_replacements)
        """
        updates_made = 0
        
        # Define specific value replacements: (old_value, new_value, value_format, context_keywords)
        # Using exact Q2 values found in the template document
        specific_replacements = [
            (values.prev_gtri, values.gtri, 'euro_k_decimal', ['GTRI', 'Gross Theoretical rental income']),
            (values.prev_vacancy_amount, values.financial_vacancy_amount, 'euro_k_decimal', ['vacancy']),
            (values.prev_vacancy_pct, values.financial_vacancy_pct, 'percent', ['vacancy']),
            (values.prev_gross_rental_income, values.gross_rental_income, 'euro_k_decimal', ['gross rental income']),
            (values.prev_gtri_ltm, values.gtri_ltm, 'euro_k_decimal', ['rent roll', 'yields']),
            (values.prev_unit_sales_proceeds, values.unit_sales_proceeds, 'euro_k_decimal', ['Unit sale proceeds', 'proceeds']),
            (values.prev_maintenance, values.maintenance_amount, 'euro_k', ['Maintenance']),
        ]
        
        # Apply specific replacements
        for old_val, new_val, fmt, keywords in specific_replacements:
            if new_val <= 0 or old_val <= 0:
                continue
            
            # Generate old and new formatted strings
            if fmt == 'euro_k_decimal':
                old_patterns = [
                    f'€{old_val:,.1f}k', f'€ {old_val:,.1f}k', f'€{old_val:,.1f} k',
                    f'€{old_val:.1f}k', f'€ {old_val:.1f}k',
                ]
                new_formatted = f'€{new_val:,.1f}k'
            elif fmt == 'euro_k':
                old_patterns = [
                    f'€{old_val:,.0f}k', f'€ {old_val:,.0f}k', f'€{old_val:,.0f} k',
                    f'€{int(old_val):,}k', f'€ {int(old_val):,}k',
                ]
                new_formatted = f'€{new_val:,.0f}k'
            elif fmt == 'percent':
                old_patterns = [f'{old_val:.1f}%', f'{old_val:.1f} %', f'{old_val}%']
                new_formatted = f'{new_val:.1f}%'
            else:
                continue
            
            # Try direct replacement first
            for old_pattern in old_patterns:
                if old_pattern in content:
                    content = content.replace(old_pattern, new_formatted)
                    updates_made += 1
                    logger.debug(f"Replaced '{old_pattern}' with '{new_formatted}'")
        
        # Define KPI contexts and their new values for context-aware replacement
        # Format: (context_keywords, new_value, value_format)
        kpi_mappings = [
            # GTRI - appears as "GTRI of the portfolio amounted to €X,XXXk"
            (['GTRI', 'Gross Theoretical Rental Income'], values.gtri, 'euro_k'),
            
            # Gross rental income - "Gross rental income was €X,XXXk"
            (['Gross rental income', 'gross rental income'], values.gross_rental_income, 'euro_k'),
            
            # Financial vacancy - "Financial vacancy was X.X%" or "vacancy of X.X%"
            (['Financial vacancy', 'vacancy'], values.financial_vacancy_pct, 'percent'),
            
            # Rent roll yields - uses LTM GTRI value
            (['rent roll', 'Rent roll', 'yields'], values.gtri_ltm, 'euro_k'),
            
            # Maintenance - "Maintenance expenses of €XXXk"
            (['Maintenance', 'maintenance', 'repair costs'], values.maintenance_amount, 'euro_k'),
            
            # Unit sales proceeds - "Unit sales proceeds of €XXXk"
            (['Unit sales proceeds', 'sale proceeds', 'proceeds'], values.unit_sales_proceeds, 'euro_k'),
            
            # CAPEX - "CAPEX of €XXXk"
            (['CAPEX', 'capex', 'capital expenditure'], values.capex_amount, 'euro_k'),
        ]
        
        for keywords, new_value, value_format in kpi_mappings:
            if new_value <= 0:
                continue
                
            # Find context regions containing the keywords
            for keyword in keywords:
                # Process keywords one at a time, starting search from beginning each time
                # This handles cases where content changes after replacement
                search_start = 0
                max_iterations = 20  # Safety limit to prevent infinite loops
                iterations = 0
                
                while iterations < max_iterations:
                    iterations += 1
                    kw_pos = content.find(keyword, search_start)
                    if kw_pos == -1:
                        break
                    
                    # Track if we made a replacement for this keyword occurrence
                    made_replacement = False
                    
                    # Look for Euro amounts or percentages - need ~1200 chars as numbers can be highly fragmented
                    search_region = content[kw_pos:kw_pos + 1500]
                    
                    if value_format == 'euro_k':
                        # Pattern to find Euro amounts (handles XML fragmentation)
                        # Matches: €X,XXX or €X.XXX or just the number parts
                        
                        # First try to find the complete value in one <w:t> tag
                        euro_pattern = r'>€([\d,\.]+)([km]?)</w:t>'
                        match = re.search(euro_pattern, search_region)
                        if match:
                            old_val = match.group(0)
                            suffix = match.group(2) or 'k'
                            new_formatted = f'>€{new_value:,.0f}{suffix}</w:t>'
                            
                            # Only replace if this is in the original content at the right location
                            full_pos = kw_pos + match.start()
                            if content[full_pos:full_pos + len(old_val)] == old_val:
                                content = content[:full_pos] + new_formatted + content[full_pos + len(old_val):]
                                updates_made += 1
                                made_replacement = True
                                logger.debug(f"Replaced '{old_val}' with '{new_formatted}' near '{keyword}'")
                        
                        # Handle fragmented case: €</w:t>...<w:t>X,XXX</w:t>
                        # Word often splits numbers like "3,200.6" into: "3," | "20" | "0" | "." | "6"
                        # Strategy: Find the full region, replace in middle section, then swap back
                        if not made_replacement:
                            # Find the Euro sign and everything up to the 'k' or 'm' suffix
                            full_value_pattern = r'(>€</w:t>)(.*?)(>k</w:t>|>m</w:t>)'
                            match = re.search(full_value_pattern, search_region, re.DOTALL | re.IGNORECASE)
                            
                            if match:
                                middle_section = match.group(2)  # Everything between € and k
                                
                                # Find all number/dot fragments in the middle section
                                frag_pattern = r'>([\d,\.]+)</w:t>'
                                fragments = list(re.finditer(frag_pattern, middle_section))
                                
                                if fragments:
                                    new_num = f'{new_value:,.0f}'
                                    
                                    # Rebuild middle section: first fragment gets new value, rest become empty
                                    new_middle = middle_section
                                    for i, frag in enumerate(fragments):
                                        old_text = f'>{frag.group(1)}</w:t>'
                                        if i == 0:
                                            # First fragment gets the new value
                                            new_middle = new_middle.replace(old_text, f'>{new_num}</w:t>', 1)
                                            logger.debug(f"Replaced fragment '{frag.group(1)}' with '{new_num}' near '{keyword}'")
                                        else:
                                            # Subsequent fragments become empty
                                            new_middle = new_middle.replace(old_text, '></w:t>', 1)
                                            logger.debug(f"Cleared fragment '{frag.group(1)}' near '{keyword}'")
                                    
                                    # Calculate absolute position and replace the middle section
                                    abs_start = kw_pos + match.start(2)
                                    abs_end = kw_pos + match.end(2)
                                    content = content[:abs_start] + new_middle + content[abs_end:]
                                    updates_made += 1
                                    made_replacement = True
                    
                    elif value_format == 'percent':
                        # Pattern to find percentages: X.X% or X%
                        # Handle both complete (>5.4%</w:t>) and fragmented (>5.</w:t>...<w:t>4</w:t>...<w:t>%</w:t>)
                        
                        # First try complete pattern
                        pct_pattern = r'>([\d,\.]+)\s*(%)</w:t>'
                        match = re.search(pct_pattern, search_region)
                        if match:
                            old_val = match.group(0)
                            new_formatted = f'>{new_value:.1f}%</w:t>'
                            
                            full_pos = kw_pos + match.start()
                            if content[full_pos:full_pos + len(old_val)] == old_val:
                                content = content[:full_pos] + new_formatted + content[full_pos + len(old_val):]
                                updates_made += 1
                                made_replacement = True
                                logger.debug(f"Replaced percentage '{old_val}' with '{new_formatted}' near '{keyword}'")
                        
                        # Try fragmented pattern: >X.</w:t>...<w:t>Y</w:t>
                        if not made_replacement:
                            # where X.Y is the percentage value (the % might be in same or separate element)
                            # Pattern: ">4.</w:t>...<w:t>1</w:t>" 
                            frag_pct_pattern = r'(>)(\d+)\.</w:t>.*?<w:t[^>]*>(\d+)</w:t>'
                            match = re.search(frag_pct_pattern, search_region, re.DOTALL)
                            if match:
                                old_int = match.group(2)
                                old_decimal = match.group(3)
                                
                                new_int = str(int(new_value))
                                new_decimal = str(int(round((new_value % 1) * 10)))  # Get first decimal digit
                                
                                logger.debug(f"Found fragmented percentage {old_int}.{old_decimal}% near '{keyword}', replacing with {new_int}.{new_decimal}%")
                                
                                int_pos = kw_pos + match.start(1)
                                old_int_full = f'>{old_int}.</w:t>'
                                new_int_full = f'>{new_int}.</w:t>'
                                
                                if content[int_pos:int_pos + len(old_int_full)] == old_int_full:
                                    content = content[:int_pos] + new_int_full + content[int_pos + len(old_int_full):]
                                    updates_made += 1
                                    made_replacement = True
                                    logger.debug(f"Replaced fragmented percentage integer '{old_int}.' with '{new_int}.' near '{keyword}'")
                                    
                                    # Now find and replace the decimal part
                                    dec_search_start = int_pos + len(new_int_full)
                                    dec_search_region = content[dec_search_start:dec_search_start + 300]
                                    dec_pattern = rf'<w:t[^>]*>({old_decimal})</w:t>'
                                    dec_match = re.search(dec_pattern, dec_search_region)
                                    
                                    if dec_match:
                                        actual_dec_pos = dec_search_start + dec_match.start(1)
                                        if content[actual_dec_pos:actual_dec_pos + len(old_decimal)] == old_decimal:
                                            content = content[:actual_dec_pos] + new_decimal + content[actual_dec_pos + len(old_decimal):]
                                            updates_made += 1
                                            logger.debug(f"Replaced fragmented percentage decimal '{old_decimal}' with '{new_decimal}' near '{keyword}'")
                    
                    # Move search_start past this keyword occurrence
                    # Always move past the current position to avoid infinite loops
                    search_start = kw_pos + len(keyword)
        
        return content, updates_made
    
    def _apply_fragmented_replacements(self, content: str, replacements: List[Tuple[str, str]]) -> str:
        """
        Apply replacements handling fragmented text in Word XML.
        
        Word often splits text across multiple <w:t> elements. This method handles:
        1. Direct text replacements (when text is not fragmented)
        2. XML-aware pattern replacements for common fragments
        """
        import re
        
        # First, apply direct replacements
        for old_text, new_text in replacements:
            if old_text in content:
                content = content.replace(old_text, new_text)
        
        # Apply common fragmented patterns using XML-aware regex
        # These handle cases where Word splits text across <w:t> elements
        content = self._apply_xml_aware_patterns(content, replacements)
        
        return content
    
    def _apply_fragmented_euro_replacements(self, content: str, values: ReportValues) -> Tuple[str, int]:
        """
        Replace fragmented euro/percent KPI values in Word XML.

        Uses template-derived previous values (``values.prev_*``) so replacements
        work after multiple reporting cycles, not only from one hardcoded anchor.
        """
        updates_made = 0

        euro_replacements = [
            (values.prev_gtri, values.gtri, lambda v: f'{v:,.1f}'),
            (values.prev_gross_rental_income, values.gross_rental_income, lambda v: f'{v:,.1f}'),
            (values.prev_vacancy_amount, values.financial_vacancy_amount, lambda v: f'{v:,.1f}'),
            (values.prev_gtri_ltm, values.gtri_ltm, lambda v: f'{v:,.1f}'),
            (values.prev_unit_sales_proceeds, values.unit_sales_proceeds, lambda v: f'{int(round(v))}'),
            (values.prev_maintenance, values.maintenance_amount, lambda v: f'{v:,.0f}'),
        ]

        for old_value, new_value, formatter in euro_replacements:
            if new_value <= 0 or old_value <= 0:
                continue
            span = _find_fragmented_number_span(content, old_value)
            if not span:
                continue
            new_text = formatter(new_value)
            new_span = _collapse_xml_runs_to_single_value(span, new_text)
            if new_span != span:
                content = content.replace(span, new_span, 1)
                updates_made += 1

        if values.financial_vacancy_pct > 0 and values.prev_vacancy_pct > 0:
            pct_span = _find_fragmented_number_span(content, values.prev_vacancy_pct)
            if pct_span:
                new_pct = f'{values.financial_vacancy_pct:.1f}'
                new_span = _collapse_xml_runs_to_single_value(pct_span, new_pct)
                if new_span != pct_span:
                    content = content.replace(pct_span, new_span, 1)
                    updates_made += 1

        return content, updates_made
    
    def _apply_fragmented_date_replacements(self, content: str, replacements: List[Tuple[str, str]]) -> str:
        """
        Apply replacements for fragmented dates in Word XML.
        
        Dates like "1-7-2025" are often fragmented as "1- 7 -202 5" across multiple
        <w:t> elements. This method handles these patterns for rent roll dates.
        
        Args:
            content: The XML content to process
            replacements: List of (old, new) text replacement tuples
            
        Returns:
            Updated content with date replacements applied
        """
        import re
        
        # Extract old and new quarter info from replacements
        old_q_num = None
        new_q_num = None
        for old_text, new_text in replacements:
            old_match = re.match(r'Q(\d)\s+(\d{4})', old_text)
            new_match = re.match(r'Q(\d)\s+(\d{4})', new_text)
            if old_match and new_match:
                old_q_num = int(old_match.group(1))
                new_q_num = int(new_match.group(1))
                break
        
        if not old_q_num:
            return content
        
        # Calculate rent roll date months (first day of month after quarter end)
        # Q2 ends June -> rent roll 1 July (7)
        # Q3 ends September -> rent roll 1 October (10)
        old_month = (old_q_num * 3) + 1
        if old_month > 12:
            old_month = 1
        new_month = (new_q_num * 3) + 1
        if new_month > 12:
            new_month = 1
        
        logger.debug(f"Replacing rent roll date month: {old_month} -> {new_month}")
        
        # Find "per 1-" context followed by month number
        # Pattern matches: >per </w:t>...<w:t>1-</w:t>...<w:t> 7 </w:t>
        # We need to replace the " 7 " with " 10 "
        
        # First, find "1-" followed by month in the XML
        # The pattern is: >1-</w:t></w:r><w:r...><w:t> 7 </w:t>
        date_pattern = rf'(>1-</w:t>)(.*?)(<w:t[^>]*>)(\s*){old_month}(\s*)(</w:t>)'
        
        def replace_month(match):
            before = match.group(1)
            middle = match.group(2)
            tag_open = match.group(3)
            space_before = match.group(4)
            space_after = match.group(5)
            tag_close = match.group(6)
            return f'{before}{middle}{tag_open}{space_before}{new_month}{space_after}{tag_close}'
        
        content = re.sub(date_pattern, replace_month, content, flags=re.DOTALL)
        
        return content
    
    def _apply_xml_aware_patterns(self, content: str, replacements: List[Tuple[str, str]]) -> str:
        """
        Apply XML-aware patterns for commonly fragmented quarter text.

        Processes every canonical ``Q{n} YYYY`` replacement pair (not just the
        calendar-previous quarter) so stale template references are updated too.
        """
        import re

        seen_pairs: set = set()
        for old_text, new_text in replacements:
            old_match = re.match(r'Q(\d)\s+(\d{4})', old_text)
            new_match = re.match(r'Q(\d)\s+(\d{4})', new_text)
            if not old_match or not new_match:
                continue
            key = (old_match.group(1), old_match.group(2))
            if key in seen_pairs:
                continue
            seen_pairs.add(key)
            content = self._apply_xml_quarter_pair(
                content,
                old_match.group(1),
                old_match.group(2),
                new_match.group(1),
                new_match.group(2),
            )
        return content

    def _apply_xml_quarter_pair(
        self,
        content: str,
        old_q_num: str,
        old_year: str,
        new_q_num: str,
        new_year: str,
    ) -> str:
        """Apply fragmented XML quarter patterns for one stale → current pair."""
        import re

        next_wt_pattern = r'</w:t></w:r><w:r[^>]*>(?:<w:rPr>(?:(?!</w:rPr>).)*</w:rPr>)?<w:t[^>]*>'

        fragmented_patterns = [
            rf'(in Q{next_wt_pattern}){old_q_num}(</w:t>)',
            rf'(% in Q{next_wt_pattern}){old_q_num}(</w:t>)',
            rf'(In Q{next_wt_pattern}){old_q_num}(</w:t>)',
            rf'( Q{next_wt_pattern}){old_q_num}(</w:t>)',
            rf'(Actions Q{next_wt_pattern}){old_q_num}(</w:t>)',
            rf'(– Q{next_wt_pattern}){old_q_num}(</w:t>)',
            rf'(- Q{next_wt_pattern}){old_q_num}(</w:t>)',
            rf'(s Q{next_wt_pattern}){old_q_num}(</w:t>)',
            rf'(>Q{next_wt_pattern}){old_q_num}(</w:t>)',
        ]
        for pattern in fragmented_patterns:
            content = re.sub(
                pattern, rf'\g<1>{new_q_num}\g<2>', content, flags=re.DOTALL
            )

        xml_patterns = [
            (rf'>Q{old_q_num}</w:t>', f'>Q{new_q_num}</w:t>'),
            (rf'>Q{old_q_num}<', f'>Q{new_q_num}<'),
            (rf'>Q{old_q_num} {old_year}<', f'>Q{new_q_num} {new_year}<'),
            (rf'> in Q{old_q_num}<', f'> in Q{new_q_num}<'),
            (rf'>in Q{old_q_num}<', f'>in Q{new_q_num}<'),
            (rf' in Q{old_q_num} {old_year}', f' in Q{new_q_num} {new_year}'),
            (rf'in Q{old_q_num} {old_year}', f'in Q{new_q_num} {new_year}'),
            (rf'>Actions Q{old_q_num} {old_year}<', f'>Actions Q{new_q_num} {new_year}<'),
            (rf'Actions Q{old_q_num} {old_year}', f'Actions Q{new_q_num} {new_year}'),
            (rf'>Actions Q{old_q_num}<', f'>Actions Q{new_q_num}<'),
            (rf'>In Q{old_q_num} {old_year}<', f'>In Q{new_q_num} {new_year}<'),
            (rf'In Q{old_q_num} {old_year}', f'In Q{new_q_num} {new_year}'),
            (rf' Q{old_q_num} {old_year}', f' Q{new_q_num} {new_year}'),
            (rf'>{old_year} – Q{old_q_num}<', f'>{new_year} – Q{new_q_num}<'),
            (rf'>{old_year} - Q{old_q_num}<', f'>{new_year} - Q{new_q_num}<'),
            (rf'>{old_year[-2:]} – Q{old_q_num}<', f'>{new_year[-2:]} – Q{new_q_num}<'),
            (rf'{old_year} – Q{old_q_num}', f'{new_year} – Q{new_q_num}'),
        ]

        for pattern, replacement in xml_patterns:
            try:
                if re.search(pattern, content):
                    content = re.sub(pattern, replacement, content)
            except re.error as e:
                logger.debug(f"Regex error for pattern {pattern}: {e}")

        return content
    
    def _get_old_date_patterns(self, previous_quarter: str) -> List[str]:
        """
        Generate date patterns that might appear on cover page from previous quarter.
        
        IMPORTANT: Patterns are sorted longest-first to prevent substring matching issues.
        e.g., "21 juli 2025" must be matched before "1 juli 2025"
        """
        patterns = []
        
        # Extract year from quarter
        parts = previous_quarter.split()
        if len(parts) == 2:
            year = int(parts[1])
            q_num = int(parts[0][1])
            
            # Calculate previous quarter end month
            month = q_num * 3
            
            # Generate various date formats for several months around report date
            # Reports are typically dated 2-3 weeks after quarter end
            for m in range(month, min(month + 3, 13)):
                for d in range(1, 32):
                    try:
                        from datetime import datetime
                        date = datetime(year, m, d)
                        # Common date formats
                        patterns.extend([
                            date.strftime('%d %B %Y'),  # 21 July 2025
                            date.strftime('%d %b %Y'),  # 21 Jul 2025
                            # Dutch formats
                            self._format_dutch_date(date),  # 21 juli 2025
                        ])
                    except ValueError:
                        continue
        
        # Remove duplicates and empty, then SORT BY LENGTH (longest first)
        # This prevents "1 juli 2025" from matching inside "21 juli 2025"
        unique_patterns = list(set(p for p in patterns if p))
        return sorted(unique_patterns, key=len, reverse=True)
    
    def _format_dutch_date(self, date) -> str:
        """Format date in Dutch style."""
        dutch_months = {
            1: 'januari', 2: 'februari', 3: 'maart', 4: 'april',
            5: 'mei', 6: 'juni', 7: 'juli', 8: 'augustus',
            9: 'september', 10: 'oktober', 11: 'november', 12: 'december'
        }
        return f"{date.day} {dutch_months[date.month]} {date.year}"
    
    def _update_cover_page_date(self, para, old_patterns: List[str], new_date: str):
        """Update cover page date to current run date (Issue 7.2)."""
        text = para.text
        for old_pattern in old_patterns:
            if old_pattern in text:
                self._replace_in_runs(para, old_pattern, new_date)
                return  # Only replace once per paragraph
    
    def _process_paragraph(self, para, values: ReportValues, 
                          old_quarter: str, new_quarter: str):
        """Process a single paragraph for replacements."""
        text = para.text
        quarter_pairs = getattr(self, '_quarter_pairs', None)
        if quarter_pairs is None:
            stale = discover_quarter_references_in_docx(
                self.template_path, exclude=new_quarter
            )
            if old_quarter not in stale:
                stale.append(old_quarter)
            quarter_pairs = build_quarter_replacement_pairs(stale, new_quarter)

        for old_pattern, new_pattern in quarter_pairs:
            if old_pattern in text:
                self._replace_in_runs(para, old_pattern, new_pattern)
                text = para.text
        
        # Placeholder replacements
        for placeholder, attr_name in self.PLACEHOLDERS.items():
            if placeholder in text:
                value = getattr(values, attr_name, '')
                
                # Format value appropriately
                if isinstance(value, float):
                    if 'pct' in attr_name or 'vacancy' in attr_name.lower():
                        formatted = f"{value:.1f}%"
                    elif 'amount' in attr_name.lower() or 'gtri' in attr_name.lower() or 'rent' in attr_name.lower():
                        formatted = f"€{value:,.1f}k"
                    else:
                        formatted = f"{value:,.1f}"
                elif isinstance(value, int):
                    formatted = str(value)
                else:
                    formatted = str(value)
                
                self._replace_in_runs(para, placeholder, formatted)
        
        # Handle REPORT_DATE placeholder
        if '{{REPORT_DATE}}' in text:
            self._replace_in_runs(para, '{{REPORT_DATE}}', values.report_date)
        
        # Handle narrative placeholders
        if '{{UNIT_SALES_NARRATIVE}}' in text and values.unit_sales_narrative:
            self._replace_in_runs(para, '{{UNIT_SALES_NARRATIVE}}', values.unit_sales_narrative)
        
        if '{{MAINTENANCE_DETAIL}}' in text and values.maintenance_detail:
            self._replace_in_runs(para, '{{MAINTENANCE_DETAIL}}', values.maintenance_detail)
        
        if '{{SUSTAINABILITY_DETAIL}}' in text and values.sustainability_detail:
            self._replace_in_runs(para, '{{SUSTAINABILITY_DETAIL}}', values.sustainability_detail)
    
    def _replace_in_runs(self, para, old_text: str, new_text: str):
        """Replace text while preserving run formatting."""
        # Simple case: text is in a single run
        for run in para.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                return
        
        # Complex case: text spans multiple runs
        # Reconstruct and replace
        full_text = para.text
        if old_text in full_text:
            new_full = full_text.replace(old_text, new_text)
            
            # Clear all runs and add new text with first run's formatting
            if para.runs:
                first_run = para.runs[0]
                for run in para.runs:
                    run.text = ''
                first_run.text = new_full

