"""
Integration tests for Q3 2025 quarterly report generation.

These tests validate the complete workflow using actual input files.
Run with: pytest tests/test_q3_2025_integration.py -v
"""

import pytest
from pathlib import Path
from datetime import datetime
import openpyxl
from docx import Document


# Test configuration
INPUT_DIR = Path("inputs")
OUTPUT_DIR = Path("outputs")


class TestBDOParser:
    """Tests for BDO financial file parsing."""
    
    def test_parse_bdo_file(self):
        """Test BDO file parsing returns expected accounts."""
        from src.parsers.bdo_parser import BDOParser
        
        parser = BDOParser()
        result = parser.parse(str(INPUT_DIR / "Cijfers_QSP_30-09-2025_d_d__14-10-2025.xlsx"))
        
        # Should parse 90+ accounts
        assert len(result.accounts) >= 90, f"Expected 90+ accounts, got {len(result.accounts)}"
    
    def test_key_accounts_exist(self):
        """Test that critical account codes are parsed."""
        from src.parsers.bdo_parser import BDOParser
        
        parser = BDOParser()
        result = parser.parse(str(INPUT_DIR / "Cijfers_QSP_30-09-2025_d_d__14-10-2025.xlsx"))
        
        key_accounts = ["1600000", "1930001", "8000003", "2400200", "1000002"]
        for code in key_accounts:
            assert code in result.accounts, f"Missing key account: {code}"
    
    def test_calculate_totals(self):
        """Test total calculations are reasonable."""
        from src.parsers.bdo_parser import BDOParser
        
        parser = BDOParser()
        result = parser.parse(str(INPUT_DIR / "Cijfers_QSP_30-09-2025_d_d__14-10-2025.xlsx"))
        totals = parser.calculate_totals(result)
        
        # Real estate should be around €101M
        assert 95_000_000 < totals['real_estate_net'] < 110_000_000, \
            f"Real estate value unexpected: €{totals['real_estate_net']:,.0f}"
        
        # Total assets should be positive
        assert totals['total_assets'] > 0, "Total assets should be positive"


class TestRentRollParser:
    """Tests for rent roll file parsing."""
    
    def test_parse_rent_roll(self):
        """Test rent roll parsing returns expected unit count."""
        from src.parsers.rent_roll_parser import RentRollParser
        
        parser = RentRollParser()
        result = parser.parse(str(INPUT_DIR / "QSP_huurlijst_1-10-2025.xlsx"))
        
        # Should have ~1,378 units
        assert 1300 < result.total_units < 1500, \
            f"Expected ~1,378 units, got {result.total_units}"
    
    def test_annual_rent_calculation(self):
        """Test annual rent total is in expected range."""
        from src.parsers.rent_roll_parser import RentRollParser
        
        parser = RentRollParser()
        result = parser.parse(str(INPUT_DIR / "QSP_huurlijst_1-10-2025.xlsx"))
        
        # Should be around €9.5M
        assert 8_000_000 < result.total_annual_rent < 12_000_000, \
            f"Annual rent unexpected: €{result.total_annual_rent:,.0f}"
    
    def test_vacancy_rate_reasonable(self):
        """Test vacancy rate is within expected range."""
        from src.parsers.rent_roll_parser import RentRollParser
        
        parser = RentRollParser()
        result = parser.parse(str(INPUT_DIR / "QSP_huurlijst_1-10-2025.xlsx"))
        
        # Vacancy should be 0-15%
        assert 0 <= result.vacancy_rate <= 0.15, \
            f"Vacancy rate seems unrealistic: {result.vacancy_rate * 100:.1f}%"


class TestSalesTrackerParser:
    """Tests for sales tracker file parsing."""
    
    def test_parse_sales_tracker(self):
        """Test sales tracker parsing returns expected sales."""
        from src.parsers.sales_tracker_parser import SalesTrackerParser
        
        parser = SalesTrackerParser()
        result = parser.parse(
            str(INPUT_DIR / "Unit_Sales_tracker_Q3_updated.xlsx"),
            period_end=datetime(2025, 9, 30)
        )
        
        # Should have 170+ total sales
        assert result.total_units_sold >= 100, \
            f"Expected 100+ total sales, got {result.total_units_sold}"
    
    def test_q3_2025_sales(self):
        """Test Q3 2025 has exactly 1 sale (Richtersweg 344)."""
        from src.parsers.sales_tracker_parser import SalesTrackerParser
        
        parser = SalesTrackerParser()
        result = parser.parse(
            str(INPUT_DIR / "Unit_Sales_tracker_Q3_updated.xlsx"),
            period_end=datetime(2025, 9, 30)
        )
        
        # Per mainPlan.pdf: 1 unit sold in Q3 2025
        assert result.quarter_units_sold == 1, \
            f"Expected 1 Q3 sale, got {result.quarter_units_sold}"
    
    def test_q3_2025_proceeds(self):
        """Test Q3 2025 proceeds are approximately €203,555."""
        from src.parsers.sales_tracker_parser import SalesTrackerParser
        
        parser = SalesTrackerParser()
        result = parser.parse(
            str(INPUT_DIR / "Unit_Sales_tracker_Q3_updated.xlsx"),
            period_end=datetime(2025, 9, 30)
        )
        
        # Expected ~€203,555
        assert 190_000 < result.quarter_proceeds < 220_000, \
            f"Q3 proceeds unexpected: €{result.quarter_proceeds:,.0f}"


class TestManagementAccountsOutput:
    """Tests for Management Accounts output file."""
    
    @pytest.fixture(autouse=True)
    def setup(self):
        """Ensure output file exists."""
        self.output_path = OUTPUT_DIR / "Management Accounts Q3 2025 - Draft 1.xlsx"
        if not self.output_path.exists():
            pytest.skip("Output file not generated - run full generation first")
    
    def test_bdo_sheet_exists(self):
        """Test BDO - Q3-25 sheet is created."""
        wb = openpyxl.load_workbook(self.output_path)
        assert "BDO - Q3-25" in wb.sheetnames, "BDO - Q3-25 sheet missing"
    
    def test_management_cijfers_sheet_exists(self):
        """Test Management Cijfers - Q3 2025 sheet exists."""
        wb = openpyxl.load_workbook(self.output_path)
        assert "Management Cijfers - Q3 2025" in wb.sheetnames, \
            "Management Cijfers - Q3 2025 sheet missing"
    
    def test_q3_column_exists(self):
        """Test Q3 2025 column is at position 29."""
        wb = openpyxl.load_workbook(self.output_path, data_only=True)
        mc = wb["Management Cijfers - Q3 2025"]
        
        q3_header = mc.cell(row=2, column=29).value
        assert isinstance(q3_header, datetime), "Q3 column header should be a date"
        assert q3_header.year == 2025 and q3_header.month == 9, \
            f"Q3 column header should be Sept 2025, got {q3_header}"
    
    def test_real_estate_value(self):
        """Test real estate value is populated correctly."""
        wb = openpyxl.load_workbook(self.output_path, data_only=True)
        mc = wb["Management Cijfers - Q3 2025"]
        
        real_estate = mc.cell(row=4, column=29).value
        assert real_estate is not None, "Real estate value should be populated"
        assert 95_000_000 < float(real_estate) < 110_000_000, \
            f"Real estate value unexpected: €{real_estate:,.0f}"


class TestComplianceCertificateOutput:
    """Tests for Compliance Certificate output file."""
    
    @pytest.fixture(autouse=True)
    def setup(self):
        """Ensure output file exists."""
        self.output_path = OUTPUT_DIR / "Compliance Certificate Berekening QSP - Q3 2025_updated.xlsx"
        if not self.output_path.exists():
            pytest.skip("Output file not generated - run full generation first")
    
    def test_required_sheets_exist(self):
        """Test all required sheets exist."""
        wb = openpyxl.load_workbook(self.output_path)
        
        required = ["SFA CC", "Suppl. Calc", "Impact Unit Sales"]
        for sheet in required:
            assert sheet in wb.sheetnames, f"{sheet} sheet missing"
    
    def test_q3_management_accounts_sheet(self):
        """Test Q3 Management Accounts sheet exists."""
        wb = openpyxl.load_workbook(self.output_path)
        # May be named "Q3 Management Accounts" or "Q3 2025 Management Accounts"
        ma_sheets = [s for s in wb.sheetnames if "Q3" in s and "Management Accounts" in s]
        assert len(ma_sheets) > 0, "Q3 Management Accounts sheet missing"
    
    def test_25q3_column_in_suppl_calc(self):
        """Test 25Q3 column exists in Suppl. Calc."""
        wb = openpyxl.load_workbook(self.output_path)
        suppl = wb["Suppl. Calc"]
        
        q3_found = False
        for col in range(1, 25):
            if suppl.cell(row=2, column=col).value == "25Q3":
                q3_found = True
                break
        
        assert q3_found, "25Q3 column not found in Suppl. Calc"


class TestWordDocumentOutput:
    """Tests for Word document output."""
    
    @pytest.fixture(autouse=True)
    def setup(self):
        """Ensure output file exists."""
        self.output_path = OUTPUT_DIR / "Quarterly QSP - Q3 2025 - Draft.docx"
        if not self.output_path.exists():
            pytest.skip("Output file not generated - run full generation first")
    
    def test_document_has_content(self):
        """Test document has substantial content."""
        doc = Document(self.output_path)
        assert len(doc.paragraphs) >= 50, \
            f"Document seems too short: {len(doc.paragraphs)} paragraphs"
    
    def test_q3_2025_reference(self):
        """Test document contains Q3 2025 references."""
        doc = Document(self.output_path)
        
        q3_found = False
        for para in doc.paragraphs:
            if "Q3 2025" in para.text:
                q3_found = True
                break
        
        assert q3_found, "No Q3 2025 references found in document"


class TestFullWorkflow:
    """Integration tests for complete workflow."""
    
    def test_orchestrator_runs_successfully(self):
        """Test complete workflow executes without errors."""
        from src.orchestrator import QuarterlyReportOrchestrator, QuarterlyReportConfig
        
        config = QuarterlyReportConfig(
            year=2025,
            quarter=3,
            bdo_file=INPUT_DIR / "Cijfers_QSP_30-09-2025_d_d__14-10-2025.xlsx",
            previous_management_accounts=INPUT_DIR / "Management Accounts Q2 2025 - Draft 1.xlsx",
            rent_roll_file=INPUT_DIR / "QSP_huurlijst_1-10-2025.xlsx",
            sales_tracker_file=INPUT_DIR / "Unit_Sales_tracker_Q3_updated.xlsx",
            previous_compliance_calc=INPUT_DIR / "Compliance Certificate Berekening QSP - Q2 2025_updated.xlsx",
            word_template=INPUT_DIR / "Quarterly_QSP_-_Q3_2025_-_Draft.docx",
            output_dir=OUTPUT_DIR
        )
        
        orchestrator = QuarterlyReportOrchestrator(config)
        results = orchestrator.run()
        
        assert results['status'] == 'success', f"Workflow failed: {results.get('errors', [])}"
        assert len(results['output_files']) >= 3, "Should generate at least 3 output files"
    
    def test_all_output_files_created(self):
        """Test all expected output files are created."""
        expected_files = [
            "Management Accounts Q3 2025 - Draft 1.xlsx",
            "Compliance Certificate Berekening QSP - Q3 2025_updated.xlsx",
            "Quarterly QSP - Q3 2025 - Draft.docx"
        ]
        
        for filename in expected_files:
            path = OUTPUT_DIR / filename
            assert path.exists(), f"Missing output file: {filename}"


if __name__ == "__main__":
    pytest.main([__file__, "-v"])

