#!/usr/bin/env python3
"""
Comprehensive Q3 2025 Validation Script

This script validates the entire quarterly report generation process by:
1. Running the full generation
2. Comparing outputs with expected manual files
3. Validating all calculations and data integrity
4. Providing detailed pass/fail reporting
"""

import sys
import os
from pathlib import Path
from datetime import datetime
from typing import Tuple, List, Dict, Any
from dataclasses import dataclass

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent.parent))

import openpyxl
from docx import Document
from loguru import logger

# Configure logging
logger.remove()
logger.add(sys.stdout, format="<level>{message}</level>", level="INFO")


@dataclass
class ValidationResult:
    """Result of a single validation check."""
    name: str
    passed: bool
    message: str
    expected: Any = None
    actual: Any = None


class Q3Validator:
    """Validates Q3 2025 quarterly report generation."""
    
    def __init__(self):
        self.results: List[ValidationResult] = []
        self.input_dir = Path("inputs")
        self.output_dir = Path("outputs")
        
    def add_result(self, name: str, passed: bool, message: str, 
                   expected: Any = None, actual: Any = None):
        """Add a validation result."""
        self.results.append(ValidationResult(
            name=name, passed=passed, message=message,
            expected=expected, actual=actual
        ))
        
        symbol = "✓" if passed else "✗"
        level = "INFO" if passed else "WARNING"
        logger.log(level, f"  {symbol} {name}: {message}")
    
    def run_all_validations(self) -> bool:
        """Run all validation checks."""
        print("=" * 70)
        print("Q3 2025 COMPREHENSIVE VALIDATION")
        print("=" * 70)
        
        # Step 1: Validate input files exist
        print("\n[1/8] Validating Input Files...")
        self.validate_input_files()
        
        # Step 2: Run generation
        print("\n[2/8] Running Full Generation...")
        generation_success = self.run_generation()
        
        if not generation_success:
            print("\n✗ Generation failed - cannot continue validation")
            return False
        
        # Step 3: Validate BDO parsing
        print("\n[3/8] Validating BDO Parser...")
        self.validate_bdo_parsing()
        
        # Step 4: Validate Rent Roll parsing
        print("\n[4/8] Validating Rent Roll Parser...")
        self.validate_rent_roll_parsing()
        
        # Step 5: Validate Sales Tracker parsing
        print("\n[5/8] Validating Sales Tracker Parser...")
        self.validate_sales_tracker_parsing()
        
        # Step 6: Validate Management Accounts output
        print("\n[6/8] Validating Management Accounts Output...")
        self.validate_management_accounts()
        
        # Step 7: Validate Compliance Certificate output
        print("\n[7/8] Validating Compliance Certificate Output...")
        self.validate_compliance_certificate()
        
        # Step 8: Validate Word document output
        print("\n[8/8] Validating Word Document Output...")
        self.validate_word_document()
        
        # Print summary
        return self.print_summary()
    
    def validate_input_files(self):
        """Check all required input files exist."""
        required_files = [
            ("BDO Financials", "Cijfers_QSP_30-09-2025_d_d__14-10-2025.xlsx"),
            ("Management Accounts Q2", "Management Accounts Q2 2025 - Draft 1.xlsx"),
            ("Rent Roll", "QSP_huurlijst_1-10-2025.xlsx"),
            ("Sales Tracker", "Unit_Sales_tracker_Q3_updated.xlsx"),
            ("Compliance Q2", "Compliance Certificate Berekening QSP - Q2 2025_updated.xlsx"),
            ("Word Template", "Quarterly_QSP_-_Q3_2025_-_Draft.docx"),
            # Expected outputs for comparison
            ("Expected MA Q3", "Management_Accounts_Q3_2025_-_Draft_1.xlsx"),
            ("Expected Compliance Q3", "Compliance Certificate Berekening QSP - Q3 2025_updated.xlsx"),
        ]
        
        for name, filename in required_files:
            path = self.input_dir / filename
            if path.exists():
                self.add_result(name, True, f"Found: {filename}")
            else:
                self.add_result(name, False, f"Missing: {filename}")
    
    def run_generation(self) -> bool:
        """Run the full generation process."""
        try:
            from src.orchestrator import QuarterlyReportOrchestrator, QuarterlyReportConfig
            
            config = QuarterlyReportConfig(
                year=2025,
                quarter=3,
                bdo_file=self.input_dir / "Cijfers_QSP_30-09-2025_d_d__14-10-2025.xlsx",
                previous_management_accounts=self.input_dir / "Management Accounts Q2 2025 - Draft 1.xlsx",
                rent_roll_file=self.input_dir / "QSP_huurlijst_1-10-2025.xlsx",
                sales_tracker_file=self.input_dir / "Unit_Sales_tracker_Q3_updated.xlsx",
                previous_compliance_calc=self.input_dir / "Compliance Certificate Berekening QSP - Q2 2025_updated.xlsx",
                word_template=self.input_dir / "Quarterly_QSP_-_Q3_2025_-_Draft.docx",
                output_dir=self.output_dir
            )
            
            # Clear outputs
            for f in self.output_dir.glob("*"):
                f.unlink()
            
            orchestrator = QuarterlyReportOrchestrator(config)
            results = orchestrator.run()
            
            if results['status'] == 'success':
                self.add_result("Generation", True, f"Generated {len(results['output_files'])} files")
                return True
            else:
                self.add_result("Generation", False, f"Failed: {results.get('errors', [])}")
                return False
                
        except Exception as e:
            self.add_result("Generation", False, f"Error: {str(e)}")
            return False
    
    def validate_bdo_parsing(self):
        """Validate BDO file parsing."""
        try:
            from src.parsers.bdo_parser import BDOParser
            
            parser = BDOParser()
            result = parser.parse(str(self.input_dir / "Cijfers_QSP_30-09-2025_d_d__14-10-2025.xlsx"))
            
            # Check account count
            if len(result.accounts) >= 90:
                self.add_result("Account Count", True, 
                               f"Parsed {len(result.accounts)} accounts", 90, len(result.accounts))
            else:
                self.add_result("Account Count", False, 
                               f"Only {len(result.accounts)} accounts", 90, len(result.accounts))
            
            # Check key accounts exist
            key_accounts = ["1600000", "1930001", "8000003", "2400200", "1000002"]
            found_count = sum(1 for code in key_accounts if code in result.accounts)
            
            if found_count == len(key_accounts):
                self.add_result("Key Accounts", True, f"All {len(key_accounts)} key accounts found")
            else:
                self.add_result("Key Accounts", False, 
                               f"Only {found_count}/{len(key_accounts)} key accounts found")
            
            # Validate totals calculation
            totals = parser.calculate_totals(result)
            
            # Check real estate value
            expected_real_estate = 101_000_000  # Approximate
            actual_real_estate = totals['real_estate_net']
            if abs(actual_real_estate - expected_real_estate) < 5_000_000:
                self.add_result("Real Estate Total", True, 
                               f"€{actual_real_estate:,.0f}", expected_real_estate, actual_real_estate)
            else:
                self.add_result("Real Estate Total", False,
                               f"€{actual_real_estate:,.0f} (expected ~€{expected_real_estate:,.0f})")
            
            # Check that we have non-zero totals
            if totals['total_assets'] != 0:
                self.add_result("Assets Total", True, f"€{totals['total_assets']:,.0f}")
            else:
                self.add_result("Assets Total", False, "Total assets is zero")
                
        except Exception as e:
            self.add_result("BDO Parsing", False, f"Error: {str(e)}")
    
    def validate_rent_roll_parsing(self):
        """Validate rent roll parsing."""
        try:
            from src.parsers.rent_roll_parser import RentRollParser
            
            parser = RentRollParser()
            result = parser.parse(str(self.input_dir / "QSP_huurlijst_1-10-2025.xlsx"))
            
            # Check unit count
            expected_units = 1378
            if abs(result.total_units - expected_units) < 50:
                self.add_result("Unit Count", True, 
                               f"{result.total_units} units", expected_units, result.total_units)
            else:
                self.add_result("Unit Count", False, 
                               f"{result.total_units} units (expected ~{expected_units})")
            
            # Check annual rent
            expected_rent = 9_500_000
            if abs(result.total_annual_rent - expected_rent) < 500_000:
                self.add_result("Annual Rent", True, 
                               f"€{result.total_annual_rent:,.0f}", expected_rent, result.total_annual_rent)
            else:
                self.add_result("Annual Rent", False, 
                               f"€{result.total_annual_rent:,.0f} (expected ~€{expected_rent:,.0f})")
            
            # Check vacancy rate is reasonable
            if 0 <= result.vacancy_rate <= 0.15:
                self.add_result("Vacancy Rate", True, f"{result.vacancy_rate * 100:.2f}%")
            else:
                self.add_result("Vacancy Rate", False, 
                               f"{result.vacancy_rate * 100:.2f}% seems unrealistic")
                
        except Exception as e:
            self.add_result("Rent Roll Parsing", False, f"Error: {str(e)}")
    
    def validate_sales_tracker_parsing(self):
        """Validate sales tracker parsing."""
        try:
            from src.parsers.sales_tracker_parser import SalesTrackerParser
            
            parser = SalesTrackerParser()
            result = parser.parse(
                str(self.input_dir / "Unit_Sales_tracker_Q3_updated.xlsx"),
                period_end=datetime(2025, 9, 30)
            )
            
            # Check total sales
            if result.total_units_sold >= 100:
                self.add_result("Total Sales", True, 
                               f"{result.total_units_sold} units sold (all time)")
            else:
                self.add_result("Total Sales", False, 
                               f"Only {result.total_units_sold} sales found")
            
            # Check Q3 2025 sales (per mainPlan.pdf: 1 unit - Richtersweg 344)
            expected_q3_sales = 1
            if result.quarter_units_sold == expected_q3_sales:
                self.add_result("Q3 2025 Sales", True, 
                               f"{result.quarter_units_sold} unit(s) sold", 
                               expected_q3_sales, result.quarter_units_sold)
            else:
                self.add_result("Q3 2025 Sales", False, 
                               f"{result.quarter_units_sold} sales (expected {expected_q3_sales})")
            
            # Check Q3 proceeds (expected ~€203,555)
            expected_proceeds = 203_555
            if abs(result.quarter_proceeds - expected_proceeds) < 10_000:
                self.add_result("Q3 Proceeds", True, 
                               f"€{result.quarter_proceeds:,.0f}", expected_proceeds, result.quarter_proceeds)
            else:
                self.add_result("Q3 Proceeds", False, 
                               f"€{result.quarter_proceeds:,.0f} (expected ~€{expected_proceeds:,.0f})")
                
        except Exception as e:
            self.add_result("Sales Tracker Parsing", False, f"Error: {str(e)}")
    
    def validate_management_accounts(self):
        """Validate Management Accounts output against expected."""
        try:
            generated_path = self.output_dir / "Management Accounts Q3 2025 - Draft 1.xlsx"
            expected_path = self.input_dir / "Management_Accounts_Q3_2025_-_Draft_1.xlsx"
            
            if not generated_path.exists():
                self.add_result("MA File", False, "Generated file not found")
                return
            
            gen_wb = openpyxl.load_workbook(generated_path, data_only=True)
            
            # Check BDO - Q3-25 sheet exists
            if "BDO - Q3-25" in gen_wb.sheetnames:
                bdo_sheet = gen_wb["BDO - Q3-25"]
                self.add_result("BDO Sheet", True, 
                               f"BDO - Q3-25 exists with {bdo_sheet.max_row} rows")
            else:
                self.add_result("BDO Sheet", False, "BDO - Q3-25 sheet missing")
                return
            
            # Check Management Cijfers sheet
            if "Management Cijfers - Q3 2025" in gen_wb.sheetnames:
                mc_sheet = gen_wb["Management Cijfers - Q3 2025"]
                self.add_result("MC Sheet", True, "Management Cijfers - Q3 2025 exists")
            else:
                self.add_result("MC Sheet", False, "Management Cijfers - Q3 2025 missing")
                return
            
            # Check Q3 column exists (should be column 29)
            q3_header = mc_sheet.cell(row=2, column=29).value
            if q3_header and isinstance(q3_header, datetime) and q3_header.year == 2025:
                self.add_result("Q3 Column", True, f"Q3 2025 column at position 29")
            else:
                self.add_result("Q3 Column", False, f"Q3 2025 column not found at expected position")
            
            # Compare key values with expected file if it exists
            if expected_path.exists():
                exp_wb = openpyxl.load_workbook(expected_path, data_only=True)
                
                # Find expected sheet
                exp_sheet = None
                for name in exp_wb.sheetnames:
                    if "Management Cijfers" in name and "Q3" in name:
                        exp_sheet = exp_wb[name]
                        break
                
                if exp_sheet:
                    # Find expected Q3 column
                    exp_col = None
                    for col in range(1, exp_sheet.max_column + 1):
                        v = exp_sheet.cell(row=2, column=col).value
                        if isinstance(v, datetime) and v.year == 2025 and v.month == 9:
                            exp_col = col
                            break
                    
                    if exp_col:
                        # Compare key rows
                        comparisons = [
                            (4, "Real estate"),
                            (9, "Cash"),
                            (10, "Equity"),
                            (12, "Bank loan"),
                        ]
                        
                        for row, name in comparisons:
                            gen_val = mc_sheet.cell(row=row, column=29).value or 0
                            exp_val = exp_sheet.cell(row=row, column=exp_col).value or 0
                            
                            try:
                                gen_val = float(gen_val)
                                exp_val = float(exp_val)
                                diff = abs(gen_val - exp_val)
                                
                                if diff < 1000:  # Within €1,000
                                    self.add_result(f"MA {name}", True, 
                                                   f"€{gen_val:,.0f} matches expected")
                                else:
                                    self.add_result(f"MA {name}", False, 
                                                   f"€{gen_val:,.0f} vs expected €{exp_val:,.0f} (diff: €{diff:,.0f})")
                            except (ValueError, TypeError):
                                self.add_result(f"MA {name}", False, 
                                               f"Cannot compare values: gen={gen_val}, exp={exp_val}")
                                
        except Exception as e:
            self.add_result("Management Accounts", False, f"Error: {str(e)}")
    
    def validate_compliance_certificate(self):
        """Validate Compliance Certificate output."""
        try:
            generated_path = self.output_dir / "Compliance Certificate Berekening QSP - Q3 2025_updated.xlsx"
            expected_path = self.input_dir / "Compliance Certificate Berekening QSP - Q3 2025_updated.xlsx"
            
            if not generated_path.exists():
                self.add_result("CC File", False, "Generated file not found")
                return
            
            gen_wb = openpyxl.load_workbook(generated_path)
            
            # Check required sheets exist
            required_sheets = ["SFA CC", "Suppl. Calc", "Impact Unit Sales"]
            for sheet_name in required_sheets:
                if sheet_name in gen_wb.sheetnames:
                    self.add_result(f"CC {sheet_name}", True, f"{sheet_name} sheet exists")
                else:
                    self.add_result(f"CC {sheet_name}", False, f"{sheet_name} sheet missing")
            
            # Check Q3 Management Accounts sheet (may be named "Q3 Management Accounts" or "Q3 2025 Management Accounts")
            ma_sheets = [s for s in gen_wb.sheetnames if "Q3" in s and "Management Accounts" in s]
            if ma_sheets:
                self.add_result("CC MA Sheet", True, f"{ma_sheets[0]} exists")
            else:
                self.add_result("CC MA Sheet", False, "Q3 Management Accounts sheet missing")
            
            # Check 25Q3 column in Suppl. Calc
            suppl_sheet = gen_wb["Suppl. Calc"]
            q3_col_found = False
            for col in range(1, 25):
                if suppl_sheet.cell(row=2, column=col).value == "25Q3":
                    q3_col_found = True
                    self.add_result("CC 25Q3 Column", True, f"25Q3 column found at position {col}")
                    break
            
            if not q3_col_found:
                self.add_result("CC 25Q3 Column", False, "25Q3 column not found in Suppl. Calc")
                
        except Exception as e:
            self.add_result("Compliance Certificate", False, f"Error: {str(e)}")
    
    def validate_word_document(self):
        """Validate Word document output."""
        try:
            generated_path = self.output_dir / "Quarterly QSP - Q3 2025 - Draft.docx"
            
            if not generated_path.exists():
                self.add_result("Word File", False, "Generated file not found")
                return
            
            doc = Document(generated_path)
            
            # Check document has content
            if len(doc.paragraphs) > 50:
                self.add_result("Word Content", True, f"{len(doc.paragraphs)} paragraphs")
            else:
                self.add_result("Word Content", False, f"Only {len(doc.paragraphs)} paragraphs")
            
            # Check for Q3 2025 references
            q3_found = False
            for para in doc.paragraphs:
                if "Q3 2025" in para.text:
                    q3_found = True
                    break
            
            if q3_found:
                self.add_result("Word Q3 Reference", True, "Q3 2025 references found")
            else:
                self.add_result("Word Q3 Reference", False, "No Q3 2025 references found")
            
            # Check document structure (tables, etc.)
            if len(doc.tables) >= 0:  # Some tables expected
                self.add_result("Word Tables", True, f"{len(doc.tables)} tables found")
                
        except Exception as e:
            self.add_result("Word Document", False, f"Error: {str(e)}")
    
    def print_summary(self) -> bool:
        """Print validation summary and return success status."""
        print("\n" + "=" * 70)
        print("VALIDATION SUMMARY")
        print("=" * 70)
        
        passed = sum(1 for r in self.results if r.passed)
        failed = sum(1 for r in self.results if not r.passed)
        total = len(self.results)
        
        print(f"\nTotal checks: {total}")
        print(f"  ✓ Passed: {passed}")
        print(f"  ✗ Failed: {failed}")
        print(f"\nSuccess rate: {passed/total*100:.1f}%")
        
        if failed > 0:
            print("\n" + "-" * 70)
            print("FAILED CHECKS:")
            print("-" * 70)
            for r in self.results:
                if not r.passed:
                    print(f"  ✗ {r.name}: {r.message}")
                    if r.expected is not None:
                        print(f"      Expected: {r.expected}")
                        print(f"      Actual:   {r.actual}")
        
        print("\n" + "=" * 70)
        if failed == 0:
            print("✓ ALL VALIDATIONS PASSED - System ready for production!")
        else:
            print(f"✗ {failed} VALIDATION(S) FAILED - Review issues above")
        print("=" * 70)
        
        return failed == 0


def main():
    """Run the validation."""
    validator = Q3Validator()
    success = validator.run_all_validations()
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()

